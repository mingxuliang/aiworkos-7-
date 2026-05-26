# -*- coding: utf-8 -*-
"""Export sandbox test scripts guide to Word (.docx)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def set_cell_shading(cell, fill: str) -> None:
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:fill"): fill,
            qn("w:val"): "clear",
        },
    )
    shading.append(shd)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_fill: str = "D9E2F3",
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], header_fill)
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_script_section(
    doc: Document,
    title: str,
    path: str,
    purpose: str,
    prerequisites: list[str],
    command: str,
    checks: list[str],
    exit_codes: str,
) -> None:
    doc.add_heading(title, level=2)
    doc.add_paragraph(f"路径：{path}")
    doc.add_paragraph(f"用途：{purpose}")
    if prerequisites:
        doc.add_paragraph("前置条件：")
        add_bullets(doc, prerequisites)
    doc.add_paragraph("运行命令：")
    add_code_block(doc, command)
    doc.add_paragraph("验证要点：")
    add_bullets(doc, checks)
    doc.add_paragraph(f"退出码：{exit_codes}")
    doc.add_paragraph()


def main() -> None:
    out_path = Path(__file__).resolve().parent.parent / "沙箱测试脚本.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = doc.add_heading("QwenPaw 工具沙箱 — 测试脚本说明", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("文档类型：测试手册  |  版本：v1.0  |  日期：2026-05-26")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph(
        "本文档汇总仓库内沙箱相关自动化测试脚本与单元测试，"
        "用于验证 Path Jail（本地沙箱）与 Docker 沙箱（Plan B）是否按预期工作，"
        "以及聊天页「沙箱环境 / 本地环境」切换是否正确传递到后端。"
    )

    doc.add_heading("一、快速开始", level=1)
    doc.add_paragraph("在仓库根目录执行（Windows PowerShell）：")
    add_code_block(
        doc,
        "cd \"C:\\Users\\92109\\Desktop\\ai-work - 沙箱版本\"\n"
        "python scripts/sandbox_smoke_test.py\n"
        "python scripts/sandbox_integration_check.py\n"
        "python scripts/sandbox_e2e_agent_test.py\n"
        "python scripts/verify_execution_environment_toggle.py\n"
        "python scripts/sandbox_console_test.py\n"
        "python scripts/sandbox_docker_check.py",
    )
    doc.add_paragraph(
        "Docker 相关测试需先启动 Docker Desktop，并构建沙箱镜像："
    )
    add_code_block(doc, ".\\scripts\\build_sandbox_image.ps1")

    doc.add_heading("二、脚本总览", level=1)
    add_table(
        doc,
        ["脚本", "层级", "依赖", "说明"],
        [
            [
                "sandbox_smoke_test.py",
                "冒烟",
                "无 Docker",
                "PathJailGuardian 阻断沙箱外路径",
            ],
            [
                "sandbox_integration_check.py",
                "集成",
                "无 Docker",
                "ToolGuardEngine + path_jail_guardian",
            ],
            [
                "sandbox_e2e_agent_test.py",
                "E2E",
                "无 Docker",
                "read_file / write_file 工具链",
            ],
            [
                "verify_execution_environment_toggle.py",
                "E2E",
                "后端 8088",
                "聊天环境切换 + AgentRequest 传参",
            ],
            [
                "sandbox_console_test.py",
                "控制台等价",
                "config.json",
                "local + docker 双后端场景",
            ],
            [
                "sandbox_docker_check.py",
                "Docker",
                "Docker + 镜像",
                "容器内 shell 执行",
            ],
            [
                "build_sandbox_image.ps1",
                "构建",
                "Docker",
                "构建 qwenpaw-sandbox:latest",
            ],
        ],
    )

    doc.add_heading("三、各脚本详细说明", level=1)

    add_script_section(
        doc,
        "3.1 路径监狱冒烟测试",
        "scripts/sandbox_smoke_test.py",
        "最小化验证：沙箱开启时，读取沙箱外系统文件应被 PathJailGuardian 拦截；"
        "读取沙箱内相对路径应放行。",
        ["设置环境变量 QWENPAW_EXECUTION_SANDBOX_ENABLED=true", "无需启动后端"],
        "python scripts/sandbox_smoke_test.py",
        [
            "blocked_findings > 0（Windows 测 C:/Windows/win.ini，Linux 测 /etc/passwd）",
            "allowed_findings = 0（workspace/ok.txt）",
            "输出 PASS: sandbox path jail smoke test",
        ],
        "0=通过，1=失败",
    )

    add_script_section(
        doc,
        "3.2 ToolGuard 集成检查",
        "scripts/sandbox_integration_check.py",
        "通过 ToolGuardEngine 调用 read_file，确认 outside 路径产生 path_jail_guardian finding，"
        "inside 路径无 finding。",
        ["QWENPAW_EXECUTION_SANDBOX_ENABLED=true", "需在 PYTHONPATH 含 src 的环境运行（pip install -e . 或 uv run）"],
        "python scripts/sandbox_integration_check.py",
        [
            "is_sandbox_enabled() 为 True",
            "outside 读取 findings[0].guardian == path_jail_guardian",
            "inside 读取无 findings",
        ],
        "0=通过，非 0=断言失败",
    )

    add_script_section(
        doc,
        "3.3 文件工具端到端测试",
        "scripts/sandbox_e2e_agent_test.py",
        "模拟 Agent 实际调用的 read_file / write_file："
        "外读拦截、内写内读成功、../ 路径穿越写拦截。",
        ["QWENPAW_EXECUTION_SANDBOX_ENABLED=true"],
        "python scripts/sandbox_e2e_agent_test.py",
        [
            "ToolGuard 拦截 outside read",
            "read_file 返回含 outside sandbox boundary",
            "write_file 在 sandbox_root 落盘",
            "read_file 能读回 hello",
            "write ../escape.txt 被拦截",
        ],
        "0=通过，1=失败",
    )

    add_script_section(
        doc,
        "3.4 聊天执行环境切换验证",
        "scripts/verify_execution_environment_toggle.py",
        "验证 Console 聊天「沙箱/本地」切换："
        "AgentRequest.execution_sandbox_enabled 正确进入 meta；"
        "override=True 时外读被拦，override=False 时外读放行；"
        "后端 /api/health 可达。",
        [
            "清除 QWENPAW_EXECUTION_SANDBOX_* 环境变量（避免覆盖 config）",
            "后端运行于 http://127.0.0.1:8088",
            "Windows 需存在 C:\\Windows\\win.ini",
        ],
        "python scripts/verify_execution_environment_toggle.py",
        [
            "PASS console payload: AgentRequest sandbox flag preserved",
            "sandbox mode read outside: BLOCKED",
            "local mode read outside: ALLOWED",
            "PASS backend health",
        ],
        "0=全部通过，1=行为或健康检查失败",
    )

    add_script_section(
        doc,
        "3.5 控制台等价沙箱测试",
        "scripts/sandbox_console_test.py",
        "写入 config.json 的 security.execution_sandbox，"
        "分别测试 backend=local 与 backend=docker；"
        "可选调用 /api/config/security/execution-sandbox/status。",
        [
            "可写 config.json",
            "Docker 用例需 Docker 可用且镜像已构建",
        ],
        "python scripts/sandbox_console_test.py",
        [
            "local::read outside blocked",
            "local::write inside ok",
            "local::traversal blocked",
            "local::shell echo ok",
            "docker::shell docker prefix / pwd /work（Docker 可用时）",
        ],
        "0=local+docker 全过，2=仅 local 过，1=local 失败",
    )

    add_script_section(
        doc,
        "3.6 Docker 沙箱集成检查",
        "scripts/sandbox_docker_check.py",
        "设置 backend=docker，在容器内执行 echo/pwd，"
        "确认输出含 [sandbox:docker 前缀与 /work 工作目录。",
        [
            "QWENPAW_EXECUTION_SANDBOX_ENABLED=true",
            "QWENPAW_EXECUTION_SANDBOX_BACKEND=docker",
            "Docker CLI 可用",
            "镜像 qwenpaw-sandbox:latest 已构建",
        ],
        "python scripts/sandbox_docker_check.py",
        [
            "STATUS: docker_available=True, docker_image_present=True",
            "SHELL OUTPUT 含 docker-sandbox-ok",
            "含 [sandbox:docker",
            "pwd 含 /work",
        ],
        "0=通过，1=失败，2=Docker 不可用（SKIP）",
    )

    add_script_section(
        doc,
        "3.7 构建沙箱 Docker 镜像",
        "scripts/build_sandbox_image.ps1",
        "从 deploy/Dockerfile.sandbox 构建 qwenpaw-sandbox:latest。",
        ["Docker Desktop 已启动"],
        ".\\scripts\\build_sandbox_image.ps1",
        ["终端输出 Built qwenpaw-sandbox:latest", "docker images 可见该镜像"],
        "PowerShell 非零退出码表示构建失败",
    )

    doc.add_heading("四、pytest 单元测试", level=1)
    doc.add_paragraph("在仓库根目录运行全部沙箱相关单元测试：")
    add_code_block(
        doc,
        "pytest tests/unit/security/test_path_jail.py "
        "tests/unit/security/test_docker_runner.py "
        "tests/unit/security/test_sandbox_status.py "
        "tests/unit/routers/test_console_placeholder.py -q",
    )

    add_table(
        doc,
        ["测试文件", "覆盖内容"],
        [
            [
                "test_path_jail.py",
                "sandbox_root 解析、路径穿越、guardian 拦截、"
                "request override、load_sandbox_settings",
            ],
            [
                "test_docker_runner.py",
                "docker run 命令组装、挂载、超时、settings 环境变量",
            ],
            [
                "test_sandbox_status.py",
                "镜像探测、get_execution_sandbox_status 健康报告",
            ],
            [
                "test_console_placeholder.py",
                "AgentRequest / dict 中 execution_sandbox_enabled 提取",
            ],
        ],
    )

    doc.add_heading("五、推荐测试顺序", level=1)
    add_bullets(
        doc,
        [
            "单元测试：pytest tests/unit/security/ …（最快，无外部依赖）",
            "冒烟：sandbox_smoke_test.py → sandbox_integration_check.py",
            "工具链：sandbox_e2e_agent_test.py",
            "启动后端后：verify_execution_environment_toggle.py",
            "配置与 API：sandbox_console_test.py",
            "Docker：build_sandbox_image.ps1 → sandbox_docker_check.py",
            "手动：Console 聊天页切换沙箱/本地，让 Agent 读取 C:\\Windows\\win.ini 对比行为",
        ],
    )

    doc.add_heading("六、常见问题", level=1)
    add_table(
        doc,
        ["现象", "可能原因", "处理"],
        [
            [
                "沙箱开启仍能读 win.ini",
                "环境变量或 UI 未传 execution_sandbox_enabled；"
                "或 override 未生效",
                "运行 verify_execution_environment_toggle.py；"
                "检查聊天请求体含 execution_sandbox_enabled",
            ],
            [
                "docker 测试 SKIP",
                "Docker 未启动或镜像未构建",
                "启动 Docker Desktop，执行 build_sandbox_image.ps1",
            ],
            [
                "integration_check import 失败",
                "未安装包或未设置 PYTHONPATH",
                "pip install -e . 或使用已配置 dev 环境",
            ],
            [
                "health check 失败",
                "后端未监听 8088",
                "启动 qwenpaw 后端后再跑 toggle 脚本",
            ],
        ],
        header_fill="FFF2CC",
    )

    doc.add_heading("附录：关键环境变量与配置", level=1)
    add_table(
        doc,
        ["项", "说明"],
        [
            ["QWENPAW_EXECUTION_SANDBOX_ENABLED", "true/false，强制覆盖是否启用沙箱"],
            ["QWENPAW_EXECUTION_SANDBOX_BACKEND", "local 或 docker"],
            ["config.security.execution_sandbox", "持久化配置：enabled、backend、docker_image 等"],
            [
                "POST /console/chat execution_sandbox_enabled",
                "单条消息级覆盖，聊天页切换即生效",
            ],
        ],
    )

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
