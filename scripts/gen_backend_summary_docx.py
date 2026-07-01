# -*- coding: utf-8 -*-
"""Generate backend change summary Word document."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

title = doc.add_heading("后端修改汇总", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("日期：2026年6月29日")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("项目：agentic-os-shaxiang-v2")
doc.add_paragraph(
    "说明：本文档汇总今日会话中涉及的后端代码修改，含具体文件路径与修改内容。"
)

doc.add_heading("一、修改概览", level=1)
doc.add_paragraph(
    "今日后端修改主要解决两类问题：\n"
    "1. 定时任务（Cron Job）消息无法投递到企业微信等第三方频道\n"
    "2. 控制台 /console/chat/stop 接口误取消企业微信等 IM 频道正在处理的任务"
)

doc.add_heading("二、文件修改详情", level=1)

# --- models.py ---
doc.add_heading("2.1 src/aiwork/app/crons/models.py", level=2)
doc.add_paragraph(
    "修改目的：修复定时任务执行时 channel 未正确传递到 AgentRequest，"
    "导致消息默认走 console 频道而非企业微信等目标频道。"
)
doc.add_paragraph("具体修改内容：")
for title, desc in [
    ("DispatchSpec.mode 默认值", '由 "stream" 改为 "final"（生成完再发送）'),
    ("JobRuntimeSpec.timeout_seconds 默认值", "由 120 改为 300（秒）"),
    ("CronJobRequest 新增字段", "channel: Optional[str] = None"),
    (
        "_validate_task_type_fields 校验器增强",
        "task_type=agent 时，自动将 dispatch.target 的 user_id、session_id "
        "以及 dispatch.channel 写入 request",
    ),
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{title}：").bold = True
    p.add_run(desc)

doc.add_paragraph("关键代码片段：")
code1 = """class DispatchSpec(BaseModel):
    type: Literal["channel"] = "channel"
    channel: str = Field(default=DEFAULT_CHANNEL)
    target: DispatchTarget
    mode: Literal["stream", "final"] = Field(default="final")
    meta: Dict[str, Any] = Field(default_factory=dict)

class JobRuntimeSpec(BaseModel):
    max_concurrency: int = Field(default=1, ge=1)
    timeout_seconds: int = Field(default=300, ge=1)
    misfire_grace_seconds: int = Field(default=60, ge=0)

class CronJobRequest(BaseModel):
    input: Optional[Any] = None
    session_id: Optional[str] = None
    user_id: Optional[str] = None
    channel: Optional[str] = None

# 在 _validate_task_type_fields 中：
self.request = self.request.model_copy(
    update={
        "user_id": target.user_id,
        "session_id": target.session_id,
        "channel": self.dispatch.channel,
    },
)"""
p = doc.add_paragraph(code1)
for run in p.runs:
    run.font.name = "Consolas"
    run.font.size = Pt(9)

# --- executor.py ---
doc.add_heading("2.2 src/aiwork/app/crons/executor.py", level=2)
doc.add_paragraph(
    "修改目的：Cron 执行 Agent 任务时，确保 stream_query 请求中的 channel "
    "始终显式设置为 dispatch.channel，防止 Runner 回退到 DEFAULT_CHANNEL（console）。"
)
doc.add_paragraph("具体修改内容：")
for title, desc in [
    ("req[\"channel\"] 显式赋值", "调用 runner.stream_query 前强制设置 channel"),
    (
        "保留 dispatch.target.user_id",
        "不覆盖为 job.owner_user_id（第三方 IM 平台 user_id 与 JWT user_id 不同）",
    ),
    (
        "ChannelManager 解析",
        "_resolve_channel_manager 按 owner_user_id 查找 per-user ChannelManager",
    ),
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{title}：").bold = True
    p.add_run(desc)

doc.add_paragraph("关键代码片段：")
code2 = """req: Dict[str, Any] = job.request.model_dump(mode="json")
req["owner_user_id"] = job.owner_user_id or ""
req["session_id"] = target_session_id or f"cron:{job.id}"
req["channel"] = job.dispatch.channel

async for event in self._runner.stream_query(req):
    await cm.send_event(
        channel=job.dispatch.channel,
        user_id=target_user_id,
        session_id=target_session_id,
        event=event,
        meta=dispatch_meta,
    )"""
p = doc.add_paragraph(code2)
for run in p.runs:
    run.font.name = "Consolas"
    run.font.size = Pt(9)

# --- console.py ---
doc.add_heading("2.3 src/aiwork/app/routers/console.py", level=2)
doc.add_paragraph(
    "修改目的：防止控制台 UI 切换/刷新会话时，POST /console/chat/stop "
    "误取消企业微信、微信、飞书等第三方频道正在处理的 AI 任务。"
)
doc.add_paragraph("涉及接口：POST /console/chat/stop")
doc.add_paragraph("函数：post_console_chat_stop")
doc.add_paragraph("具体修改内容：")
for title, desc in [
    ("频道校验守卫", "停止前先查询 chat 的 channel 字段"),
    (
        "仅允许停止 console 频道",
        '非 console 频道返回 {"stopped": false, "reason": "..."}',
    ),
    ("session_id 解析优化", "支持 chat_id 为 session_id 时先解析为 UUID"),
    ("文档字符串更新", "明确说明不停止第三方 IM 频道任务"),
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{title}：").bold = True
    p.add_run(desc)

doc.add_paragraph("关键代码片段：")
code3 = """_CONSOLE_ONLY_CHANNELS = {"console", ""}

existing = await chat_manager.get_chat(chat_id)
if existing is not None:
    chat_channel = getattr(existing, "channel", "console") or "console"
    if chat_channel not in _CONSOLE_ONLY_CHANNELS:
        return {
            "stopped": False,
            "reason": f"channel '{chat_channel}' cannot be stopped from console",
        }

stopped = await workspace.task_tracker.request_stop(resolved_id)
return {"stopped": stopped}"""
p = doc.add_paragraph(code3)
for run in p.runs:
    run.font.name = "Consolas"
    run.font.size = Pt(9)

# --- API table ---
doc.add_heading("三、接口影响说明", level=1)
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "接口"
hdr[1].text = "方法"
hdr[2].text = "变更类型"
hdr[3].text = "说明"
for row in [
    (
        "/console/chat/stop",
        "POST",
        "行为变更",
        "新增频道校验，wecom/wechat/feishu/dingtalk 任务不会被此接口停止",
    ),
    (
        "/cron/jobs",
        "POST/PUT",
        "间接修复",
        "创建/更新任务时 dispatch.channel 会正确写入 request.channel",
    ),
    (
        "Cron 执行（内部）",
        "-",
        "Bug 修复",
        "Agent 任务执行时 channel 不再回退到 console",
    ),
]:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

doc.add_heading("四、问题与修复对应关系", level=1)
for problem, fix in [
    ("企业微信定时任务消息收不到", "crons/models.py + executor.py：channel 传播与显式赋值"),
    ("企业微信发消息后 AI 无回复（被中断）", "console.py：/console/chat/stop 频道守卫"),
    ("定时任务默认超时过短", "crons/models.py：timeout_seconds 120 → 300"),
    ("定时任务默认发送模式", "crons/models.py：dispatch.mode stream → final"),
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"问题：{problem}\n").bold = True
    p.add_run(f"修复：{fix}")

doc.add_heading("五、修改文件清单", level=1)
for f in [
    "src/aiwork/app/crons/models.py",
    "src/aiwork/app/crons/executor.py",
    "src/aiwork/app/routers/console.py",
]:
    doc.add_paragraph(f, style="List Number")

doc.add_paragraph(
    "注：频道配置保存（PUT /config/channels/{channel}）本身已有 restart_channel 逻辑，"
    "今日未修改该接口。"
)

out_path = os.path.join(
    os.path.dirname(os.path.dirname(__file__)),
    "后端修改汇总_2026-06-29.docx",
)
doc.save(out_path)
print(out_path)
