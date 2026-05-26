# -*- coding: utf-8 -*-
"""Export user isolation development plan to Word (.docx)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


def set_cell_shading(cell, fill: str) -> None:
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:fill"): fill,
            qn("w:val"): "clear",
        },
    )
    shading.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "D9E2F3") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], header_fill)
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    out_path = Path(__file__).resolve().parent.parent / "用户隔离开发方案.docx"

    doc = Document()

    # Default font for Chinese
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = doc.add_heading("QwenPaw / AI Work OS 用户隔离开发方案", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("文档类型：技术方案  |  版本：v1.0  |  日期：2026-05-26")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph(
        "本文档基于当前代码库（src/qwenpaw + console）现状分析，"
        "针对共享 Agent、用户记忆与文件未隔离问题，给出分阶段开发方案。"
    )

    # 1
    doc.add_heading("一、现状诊断", level=1)
    doc.add_paragraph(
        "系统隔离分为两层，且实现不一致：Agent 级（workspace、MEMORY.md、media 等）"
        "与 User 级（chats.json、session JSON）。"
    )

    doc.add_heading("1.1 当前隔离状态", level=2)
    add_table(
        doc,
        ["数据", "当前隔离粒度", "权限校验", "主要问题"],
        [
            ["Agent 元数据 CRUD", "Agent + user_id", "有", "仅管理接口受保护"],
            ["Agent 运行时（聊天/文件/配置/MCP）", "仅 Agent", "无", "知道 agent_id 即可访问"],
            ["对话列表 / 会话 JSON", "Agent + User", "部分", "batch-delete 等未校验"],
            ["ReMe 长期记忆", "仅 Agent", "无", "多用户共享 MEMORY.md"],
            ["Workspace 文件 / 上传", "仅 Agent", "无", "多用户共享"],
            ["前端 user_id", "—", "—", "大量写死 default"],
        ],
    )

    doc.add_heading("1.2 身份不一致问题", level=2)
    add_bullets(
        doc,
        [
            "JWT middleware 使用数字 user id；Console 聊天使用 username；前端默认 default。",
            "Agent 创建时绑定的 user_id 格式不统一，导致会话列表为空、403 等异常。",
            "user_id = null 表示共享 Agent，是产品设计而非 bug；多租户需区分共享与个人 Agent。",
        ],
    )

    doc.add_heading("1.3 存储路径现状", level=2)
    add_table(
        doc,
        ["资源", "路径模式", "隔离范围"],
        [
            ["Agent workspace", "workspaces/{agent_id}/", "Per-agent"],
            ["Chat 注册表", "{workspace}/chats.json", "Per-agent，按 user_id 过滤"],
            ["Session JSON", "{workspace}/sessions/{user_id}_{session_id}.json", "Per-agent + Per-user"],
            ["ReMe 记忆", "{workspace}/MEMORY.md, memory/*.md", "Per-agent（多用户共享）"],
            ["Light context", "{workspace}/dialog/*.jsonl", "Per-agent（多用户共享）"],
            ["上传文件", "{workspace}/media/", "Per-agent（多用户共享）"],
        ],
    )

    # 2
    doc.add_heading("二、产品策略（必须先定）", level=1)
    doc.add_paragraph("开发前需明确三种 Agent 类型及隔离策略：")
    add_table(
        doc,
        ["类型", "user_id", "记忆", "文件", "适用场景"],
        [
            ["个人 Agent", "绑定创建者", "按用户隔离", "按用户隔离", "默认模式，多租户"],
            ["团队 Agent", "绑定团队/部门", "可选共享或按用户", "共享 workspace + 用户子目录", "协作助手"],
            ["系统共享 Agent", "null", "全员共享", "全员共享", "官方模板、公共 QA"],
        ],
    )
    doc.add_paragraph(
        "推荐默认策略：新创建的 Agent 一律绑定 user_id，不再默认共享；"
        "default Agent 仅开发环境保留，生产需管理员显式标记为系统共享。"
    )

    # 3
    doc.add_heading("三、目标架构", level=1)

    doc.add_heading("3.1 统一身份（Phase 1）", level=2)
    add_bullets(
        doc,
        [
            "全系统只认一种 canonical user_id（建议 JWT 数字 id，username 仅展示）。",
            "前端登录后从 token 写入 currentUserId，禁止客户端随意传 default。",
            "后端所有写接口强制覆盖 client 传来的 user_id，不信任前端。",
        ],
    )

    doc.add_heading("3.2 访问控制（Phase 2）", level=2)
    doc.add_paragraph(
        "在 get_agent_for_request() 统一加归属校验（admin 可 bypass）："
    )
    add_bullets(
        doc,
        [
            "个人 Agent：仅 owner + admin 可访问。",
            "共享 Agent：所有已登录用户可读；写 workspace / 改配置需 admin 或显式授权。",
            "影响范围：workspace、console chat、config、MCP、skills、crons、plan 等所有 agent-scoped 路由。",
        ],
    )

    doc.add_heading("3.3 数据隔离方案对比", level=2)

    doc.add_heading("方案 A：最小改动（1–2 天）", level=3)
    add_bullets(
        doc,
        [
            "不改存储结构，仅补齐 ownership 校验 + 统一 user_id。",
            "共享 Agent 仍共享记忆（需在文档/UI 中明确说明）。",
            "适合：短期上线多用户，接受共享 Agent = 共享记忆。",
        ],
    )

    doc.add_heading("方案 B：推荐 — 记忆/文件按用户子目录（3–5 天）", level=3)
    doc.add_paragraph("目录结构：")
    code = doc.add_paragraph()
    code_run = code.add_run(
        "workspaces/{agent_id}/\n"
        "├── agent.json, AGENTS.md, SOUL.md   # 共享 persona\n"
        "├── users/{user_id}/\n"
        "│   ├── MEMORY.md, memory/\n"
        "│   ├── dialog/\n"
        "│   ├── media/\n"
        "│   └── sessions/\n"
        "└── chats.json                         # 仍按 user_id 过滤"
    )
    code_run.font.name = "Consolas"
    code_run.font.size = Pt(9)

    add_bullets(
        doc,
        [
            "ReMeLightMemoryManager：working_dir → workspace/users/{user_id}/",
            "LightContextManager：dialog_path 同上。",
            "上传/工具读写默认限制在用户子目录（Tool Guard 配合）。",
            "共享 Agent（user_id=null）仍走 agent 根目录，保持现有行为。",
        ],
    )

    doc.add_heading("方案 C：最强 — 每用户独立 Agent 实例（5–8 天）", level=3)
    add_bullets(
        doc,
        [
            "用户注册时 fork 模板 Agent → agent_id = tpl_{user_id}。",
            "完全物理隔离，适合高安全场景；运维成本高，模板更新需同步机制。",
        ],
    )
    doc.add_paragraph("建议采用方案 B，在共享 Agent 与多租户之间平衡最好。")

    # 4
    doc.add_heading("四、分阶段开发计划", level=1)

    doc.add_heading("Phase 1：身份统一（P0，约 1 天）", level=2)
    add_table(
        doc,
        ["任务", "说明"],
        [
            ["统一 JWT 解析", "console.py / api.py / agents.py 共用 _get_canonical_user_id()"],
            ["前端绑定", "登录后设置 currentUserId；Chat 不再 fallback default"],
            ["数据迁移脚本", "chats.json、session 文件名统一为 canonical id"],
            ["Agent 创建", "JWT 模式下强制 user_id=current_user"],
        ],
    )
    doc.add_paragraph("验收标准：两个用户登录后，互相看不到对方会话列表。")

    doc.add_heading("Phase 2：权限补齐（P0，约 1–2 天）", level=2)
    add_table(
        doc,
        ["任务", "说明"],
        [
            ["get_agent_for_request()", "加 _check_agent_ownership"],
            ["Chat API", "POST /chats 强制 JWT user；batch-delete 逐条校验"],
            ["console/chat/stop", "校验 chat 归属"],
            ["files/preview", "限制在 agent workspace 内"],
            ["Proactive memory", "list_chats(user_id=...) 过滤"],
        ],
    )
    doc.add_paragraph("验收标准：用户 A 无法通过 X-Agent-Id 访问用户 B 的 Agent workspace。")

    doc.add_heading("Phase 3：记忆与文件隔离（P1，约 2–3 天，方案 B）", level=2)
    add_table(
        doc,
        ["任务", "说明"],
        [
            ["UserWorkspaceResolver", "根据 (agent_id, user_id, agent.user_id) 解析实际路径"],
            ["ReMe / LightContext", "传入用户级 working_dir"],
            ["Workspace API", "列表/读写默认 scoped 到用户子目录"],
            ["Tool Guard", "工具写文件不能跨 users/{other}/"],
            ["迁移", "现有 MEMORY.md 按 owner 或 shared 归档"],
        ],
    )
    doc.add_paragraph("验收标准：同一 Agent 下，A/B 对话后各自 MEMORY.md 互不影响。")

    doc.add_heading("Phase 4：共享 Agent 产品化（P2，约 1–2 天）", level=2)
    add_table(
        doc,
        ["任务", "说明"],
        [
            ["Agent 类型字段", "visibility: private | shared | system"],
            ["UI 标识", "共享 Agent 显示「公共记忆」徽章"],
            ["权限矩阵", "shared：所有人可聊；改配置仅 admin"],
            ["默认 Agent 策略", "生产环境禁止匿名 default 命名空间"],
        ],
    )

    doc.add_heading("Phase 5：沙箱与工具隔离（P2，与 sandbox 规划衔接）", level=2)
    add_bullets(
        doc,
        [
            "Tool 执行：每次调用独立进程/容器（已有 sandbox 规划）。",
            "文件边界：users/{user_id}/ 作为 sandbox root。",
            "与 Phase 3 目录策略一致，避免重复设计。",
        ],
    )

    # 5
    doc.add_heading("五、API / 模型变更摘要", level=1)
    doc.add_paragraph("AgentProfileRef 扩展（示意）：")
    code2 = doc.add_paragraph()
    code2_run = code2.add_run(
        "class AgentProfileRef:\n"
        "    user_id: str | None          # owner；null = 系统共享\n"
        '    visibility: Literal["private", "shared", "system"] = "private"\n'
        '    memory_scope: Literal["agent", "user"] = "user"  # 新增'
    )
    code2_run.font.name = "Consolas"
    code2_run.font.size = Pt(9)

    doc.add_paragraph("建议新增内部服务 UserContextResolver：")
    add_bullets(
        doc,
        [
            "resolve_user_id(request) -> str",
            "resolve_agent_access(request, agent_id) -> AgentRef",
            "resolve_user_workspace(agent_id, user_id) -> Path",
        ],
    )
    doc.add_paragraph("所有 agent-scoped 路由只通过该 resolver，避免各处重复逻辑。")

    # 6
    doc.add_heading("六、风险与兼容", level=1)
    add_table(
        doc,
        ["风险", "应对措施"],
        [
            ["历史数据在 default 下", "一次性 migration + 管理员工具合并/分配"],
            ["渠道（微信等）user_id 格式不同", "渠道层映射到 canonical id"],
            ["共享 Agent 用户期望完全私有", "UI + 文档明确；创建时默认 private"],
            ["性能（每用户子目录）", "仅活跃用户 lazy 创建目录"],
        ],
    )

    # 7
    doc.add_heading("七、优先级与时间线", level=1)
    add_table(
        doc,
        ["阶段", "内容", "优先级"],
        [
            ["Week 1", "Phase 1 + Phase 2：堵住越权，统一身份", "必须做"],
            ["Week 2", "Phase 3：记忆/文件按用户隔离", "多租户核心"],
            ["Week 3", "Phase 4 + 5：共享 Agent 产品化 + 沙箱衔接", "增强"],
        ],
    )

    # 8
    doc.add_heading("八、结论", level=1)
    add_table(
        doc,
        ["问题", "根因", "开发方向"],
        [
            ["共享 Agent", "user_id=null 是设计", "产品分层 + 默认 private"],
            ["记忆不隔离", "ReMe 路径在 agent 根", "方案 B：users/{user_id}/memory"],
            ["文件不隔离", "workspace 无 user 子目录", "同上 + Tool Guard"],
            ["越权访问", "runtime API 无 ownership", "Phase 2 统一网关校验"],
            ['前端 "default"', "未接 JWT", "Phase 1 前端 + 后端强制覆盖"],
        ],
    )

    doc.add_paragraph()
    conclusion = doc.add_paragraph()
    conclusion.add_run("推荐路径：").bold = True
    conclusion.add_run("先完成 Phase 1 + 2（安全底线），再实施方案 B 的 Phase 3（记忆/文件隔离）。")

    doc.save(out_path)
    print(f"Exported: {out_path}")


if __name__ == "__main__":
    main()
