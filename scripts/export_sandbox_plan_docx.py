# -*- coding: utf-8 -*-
"""Export tool sandbox development plan to Word (.docx)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


def set_cell_shading(cell, fill: str) -> None:
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:fill"): fill,
            qn("w:val"): "clear",
        },
    )
    shading.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "D9E2F3") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], header_fill)
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def main() -> None:
    out_path = Path(__file__).resolve().parent.parent / "工具沙箱开发方案.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = doc.add_heading("QwenPaw / AI Work OS 工具沙箱开发方案", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("文档类型：技术方案  |  版本：v1.0  |  日期：2026-05-26")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph(
        "本文档基于当前代码库（src/qwenpaw）与产品文档现状分析，"
        "针对工具执行沙箱（对比文档标注为「规划中」）给出分阶段开发方案。"
        "与用户隔离方案配套，建议在 users/{user_id}/ 路径策略落地后实施沙箱执行层。"
    )

    # 1
    doc.add_heading("一、现状诊断", level=1)

    doc.add_heading("1.1 已有能力：策略层安全（非沙箱）", level=2)
    doc.add_paragraph(
        "当前安全栈是「检测 + 审批」，不是「隔离执行」。"
        "Tool Guard、File Guard、Skill Scanner 与 Approval Level 构成策略层，"
        "工具实际仍在 QwenPaw 同 OS 用户权限下执行。"
    )
    add_table(
        doc,
        ["组件", "作用", "是否沙箱"],
        [
            ["Tool Guard", "正则规则、Shell 逃逸检测、敏感路径拦截", "否（策略）"],
            ["File Guard", "敏感目录 deny-list（.secret 等）", "否（策略）"],
            ["Skill Scanner", "安装/启用前静态扫描", "否（安装期）"],
            ["Approval Level", "STRICT / SMART / AUTO / OFF", "否（人工门控）"],
            ["Workspace Context", "cwd、相对路径默认在 agent workspace", "部分（约定，非 OS 强制）"],
        ],
    )

    doc.add_heading("1.2 实际执行方式", level=2)
    add_table(
        doc,
        ["工具类型", "执行方式", "风险"],
        [
            ["execute_shell_command", "宿主机 subprocess，cwd=workspace", "审批通过后仍用宿主权限"],
            ["文件工具", "Python 直接 open()，绝对路径不受 workspace 限制", "可读写 workspace 外路径"],
            ["MCP", "启动 MCP server 子进程，继承宿主环境", "无独立挂载/网络策略"],
            ["Skills", "进程内加载执行（SECURITY.md 明确）", "等同于信任任意 Python 代码"],
            ["Browser", "Playwright 子进程", "浏览器级，不是 Agent 工具沙箱"],
        ],
    )

    doc.add_heading("1.3 核心缺口", level=2)
    add_table(
        doc,
        ["规划能力", "当前状态"],
        [
            ["每次工具调用独立进程/容器", "未实现"],
            ["users/{user_id}/ 作为沙箱根", "未实现（依赖用户隔离 Phase 3）"],
            ["workspace 外路径硬阻断", "仅 rm 检测 + 警告，非强制"],
            ["MCP/Skill 执行隔离", "未实现"],
            ["网络 egress 控制", "未实现"],
        ],
    )
    doc.add_paragraph(
        "结论：现有 Tool Guard 是「门卫」，沙箱是「牢房」。两者互补，不能互相替代。"
    )

    # 2
    doc.add_heading("二、设计目标", level=1)

    doc.add_heading("2.1 安全目标", level=2)
    add_bullets(
        doc,
        [
            "最小权限：工具只能访问被授权的目录、网络、环境变量。",
            "失败安全：沙箱不可用时可配置为「拒绝执行」而非静默降级。",
            "可审计：每次沙箱执行记录 tool、user、agent、挂载、退出码。",
            "多租户就绪：沙箱根与用户隔离目录一致。",
        ],
    )

    doc.add_heading("2.2 非目标（首版不做）", level=2)
    add_bullets(
        doc,
        [
            "完整 VM 级隔离。",
            "替换 Tool Guard / Approval（保留，沙箱在其之后）。",
            "所有工具 100% 容器化（分阶段、分 backend）。",
        ],
    )

    doc.add_heading("2.3 沙箱粒度（推荐）", level=2)
    add_table(
        doc,
        ["维度", "推荐策略", "说明"],
        [
            ["文件边界", "Per Agent + Per User", "workspaces/{agent_id}/users/{user_id}/ 为读写根"],
            ["进程隔离", "Per Tool Call", "每次 shell/MCP/高风险工具独立子进程"],
            ["会话级容器", "可选（Phase 3+）", "长任务复用容器，降低启动开销"],
            ["共享 Agent", "Agent 根目录只读挂载", "persona/skills 共享，用户数据写 user 子目录"],
        ],
    )

    # 3
    doc.add_heading("三、目标架构", level=1)

    doc.add_heading("3.1 执行流水线（目标态）", level=2)
    doc.add_paragraph(
        "Tool Call → Tool Guard + Approval → SandboxRunner → Backend（local/docker/firejail）"
        "→ Mount sandbox_root → Tool Executor → 结果回传 Agent。"
    )
    doc.add_paragraph(
        "最佳插入点：ToolGuardMixin._acting() 审批通过后、调用 super()._acting() 之前"
        "（tool_guard_mixin.py 约第 466 行），统一走 SandboxRunner.execute(tool_call, context)。"
    )

    doc.add_heading("3.2 沙箱上下文 SandboxContext", level=2)
    add_code_block(
        doc,
        "SandboxContext(\n"
        "    agent_id, user_id, session_id, tool_name,\n"
        "    sandbox_root,      # users/{user_id}/ 或 workspace\n"
        "    workspace_ro,      # agent 共享只读区\n"
        "    network_policy,    # deny | allowlist | full\n"
        "    env_allowlist, timeout_seconds,\n"
        "    backend,           # off | local | docker | firejail\n"
        ")",
    )
    doc.add_paragraph(
        "由 UserWorkspaceResolver（用户隔离 Phase 3）解析 sandbox_root，"
        "沙箱与用户隔离共用同一套路径策略。"
    )

    doc.add_heading("3.3 三层隔离模型", level=2)
    add_table(
        doc,
        ["层级", "机制", "负责模块"],
        [
            ["L1 策略", "Tool Guard / File Guard / Approval", "已有"],
            ["L2 路径", "强制 sandbox_root，拒绝越界绝对路径", "新 PathJailGuardian + file_io"],
            ["L3 执行", "独立进程/容器、降权、网络隔离", "新 SandboxRunner"],
        ],
    )

    # 4
    doc.add_heading("四、Backend 方案对比", level=1)

    doc.add_heading("方案 A：Local Process Jail（MVP，2–3 天）", level=2)
    doc.add_paragraph("适用：Windows / Linux 开发环境、无 Docker。")
    add_bullets(
        doc,
        [
            "Shell：子进程 + cwd=sandbox_root + 环境变量白名单。",
            "文件工具：_resolve_file_path 强制落在 sandbox_root 内。",
            "越界路径：硬拒绝（403），不只警告。",
            "Linux 增强：可选 firejail / bubblewrap。",
        ],
    )
    doc.add_paragraph("局限：仍共享 QwenPaw OS 用户，防君子不防内核级攻击。")

    doc.add_heading("方案 B：Docker Per-Call（推荐生产，5–7 天）", level=2)
    doc.add_paragraph("适用：Linux 服务器、Docker 已部署。")
    add_bullets(
        doc,
        [
            "镜像：精简 qwenpaw-sandbox:latest（python + 常用 CLI）。",
            "挂载：-v sandbox_root:/work:rw + -v workspace_ro:/ro:ro。",
            "网络：默认 --network none，allowlist 域名需 proxy。",
            "资源：--memory / --cpus / --pids-limit。",
            "MCP：MCP server 在容器内启动，stdio 转发。",
        ],
    )
    doc.add_paragraph("局限：冷启动 200–800ms；Windows 需 WSL2/Docker Desktop。")

    doc.add_heading("方案 C：Session Container（进阶，+3–5 天）", level=2)
    add_bullets(
        doc,
        [
            "会话开始时创建容器，会话内多次 tool call 复用。",
            "适合 Browser、长 Shell 流水线。",
            "需生命周期管理（idle 超时销毁、/stop 清理）。",
        ],
    )

    doc.add_paragraph(
        "推荐路径：Phase 1–2 方案 A → Phase 3–4 方案 B → Phase 5 方案 C。"
    )

    # 5
    doc.add_heading("五、分阶段开发计划", level=1)

    doc.add_heading("Phase 0：前置 — 用户隔离路径（依赖，约 2–3 天）", level=2)
    add_code_block(
        doc,
        "workspaces/{agent_id}/users/{user_id}/  → sandbox_root (RW)\n"
        "workspaces/{agent_id}/                  → workspace_ro (共享只读)",
    )
    doc.add_paragraph("验收：两个用户的 sandbox_root 物理分离。")

    doc.add_heading("Phase 1：路径强制（P0，约 2 天）", level=2)
    add_table(
        doc,
        ["任务", "说明"],
        [
            ["PathJailGuardian", "file/shell 工具路径必须在 sandbox_root 内"],
            ["file_io._resolve_file_path", "绝对路径 resolve 后校验 relative_to(sandbox_root)"],
            ["rule_guardian rm 检测", "outside workspace 从 WARN → BLOCK（可配置）"],
            ["SandboxContext resolver", "从 request context 解析 agent/user/session"],
        ],
    )
    doc.add_paragraph("验收：Agent 无法通过 read_file(\"/etc/passwd\") 读取系统文件。")

    doc.add_heading("Phase 2：SandboxRunner 框架（P0，约 2–3 天）", level=2)
    add_table(
        doc,
        ["任务", "说明"],
        [
            ["security/sandbox/runner.py", "抽象 SandboxRunner.execute()"],
            ["LocalSandboxRunner", "包装 shell/file 执行"],
            ["接入 ToolGuardMixin", "审批通过后走 Runner"],
            ["配置项", "security.execution_sandbox.enabled/backend/fail_closed"],
            ["审计日志", "结构化记录每次沙箱执行"],
        ],
    )
    add_code_block(
        doc,
        '{\n'
        '  "security": {\n'
        '    "execution_sandbox": {\n'
        '      "enabled": true,\n'
        '      "backend": "local",\n'
        '      "fail_closed": true,\n'
        '      "default_network": "deny"\n'
        '    }\n'
        '  }\n'
        '}',
    )
    doc.add_paragraph("验收：开启沙箱后 shell 在 sandbox_root 内执行；关闭时行为与现网一致。")

    doc.add_heading("Phase 3：Docker Backend（P1，约 4–5 天）", level=2)
    add_table(
        doc,
        ["任务", "说明"],
        [
            ["DockerSandboxRunner", "docker run --rm 单次调用"],
            ["沙箱镜像构建", "docker/sandbox/Dockerfile + CI"],
            ["Shell 代理", "容器内执行，stdout/stderr 回传"],
            ["MCP 沙箱化", "StdIO MCP 在容器内 spawn"],
            ["网络策略", "默认 none；HTTP 工具走 allowlist proxy"],
        ],
    )
    doc.add_paragraph("验收：容器内 rm -rf / 不影响宿主；容器无法访问其他 user 目录。")

    doc.add_heading("Phase 4：Skill 执行隔离（P1，约 3–4 天）", level=2)
    add_table(
        doc,
        ["任务", "说明"],
        [
            ["Skill 脚本分类", "标记 requires_sandbox: true"],
            ["执行器改造", "skill 脚本走 SandboxRunner，不再 in-process"],
            ["Scanner 联动", "高风险 skill 强制 sandbox"],
        ],
    )
    doc.add_paragraph("验收：恶意 skill 脚本无法读写 sandbox_root 外文件。")

    doc.add_heading("Phase 5：Session 级优化 + 控制台（P2，约 3 天）", level=2)
    add_table(
        doc,
        ["任务", "说明"],
        [
            ["Session Container 池", "复用容器，降低延迟"],
            ["控制台 Security 页", "沙箱 backend 状态、审计"],
            ["Agent 级 override", "agent.json 可设 sandbox_backend"],
            ["与 approval_level 联动", "STRICT 模式强制 docker"],
        ],
    )

    # 6
    doc.add_heading("六、与现有组件的关系", level=1)
    add_table(
        doc,
        ["场景", "Tool Guard", "沙箱"],
        [
            ["rm -rf ~/", "规则命中 → 审批", "即使批准，也只能删 sandbox_root 内"],
            ["读 /etc/passwd", "可能未命中规则", "Path Jail 硬拒绝"],
            ["已审批的 curl 外联", "允许", "网络 policy 控制"],
            ["Skill 后门", "Scanner 可能 warn", "容器内无法逃逸"],
        ],
    )

    # 7
    doc.add_heading("七、配置与 API 变更", level=1)

    doc.add_heading("7.1 SecurityConfig 扩展", level=2)
    add_code_block(
        doc,
        "class ExecutionSandboxConfig(BaseModel):\n"
        "    enabled: bool = False\n"
        '    backend: Literal["off", "local", "docker", "firejail"] = "local"\n'
        "    fail_closed: bool = True\n"
        '    network: Literal["deny", "allowlist", "full"] = "deny"\n'
        "    network_allowlist: list[str] = []\n"
        "    mount_workspace_readonly: bool = True\n"
        "    docker_image: str = \"qwenpaw-sandbox:latest\"",
    )

    doc.add_heading("7.2 Agent 级配置", level=2)
    add_code_block(
        doc,
        '{\n'
        '  "approval_level": "SMART",\n'
        '  "sandbox": {\n'
        '    "backend": "docker",\n'
        '    "network": "allowlist",\n'
        '    "network_allowlist": ["api.example.com"]\n'
        '  }\n'
        '}',
    )

    doc.add_heading("7.3 控制台", level=2)
    add_bullets(
        doc,
        [
            "Settings → Security：沙箱开关、backend 选择、fail_closed。",
            "Agent Config：覆盖全局沙箱策略。",
            "Debug 日志：沙箱执行 audit 流。",
        ],
    )

    # 8
    doc.add_heading("八、风险与兼容", level=1)
    add_table(
        doc,
        ["风险", "应对"],
        [
            ["Docker 不可用", "dev 可降级 local；生产建议 fail_closed"],
            ["性能（容器冷启动）", "Phase 5 Session 池；低风险工具走 local"],
            ["Windows 开发", "Phase 1–2 local；Docker 依赖 WSL2"],
            ["Browser/Playwright", "单独策略，不强行塞入同一容器"],
            ["ACP bypass 冲突", "bypass 模式显式跳过沙箱，日志告警"],
            ["历史工具依赖绝对路径", "迁移期 WARN + 文档"],
        ],
    )

    # 9
    doc.add_heading("九、时间线建议", level=1)
    add_table(
        doc,
        ["周期", "内容", "产出"],
        [
            ["Week 1", "Phase 0 用户路径 + Phase 1 Path Jail", "路径无法越界"],
            ["Week 2", "Phase 2 Local SandboxRunner + Mixin 接入", "统一执行网关"],
            ["Week 3–4", "Phase 3 Docker Backend", "生产级隔离"],
            ["Week 5", "Phase 4 Skill 沙箱", "扩展面收敛"],
            ["Week 6", "Phase 5 优化 + 控制台", "可运维"],
        ],
    )
    doc.add_paragraph(
        "MVP（约 1 周）：Phase 0 路径 + Phase 1 Path Jail + Phase 2 Local Runner。"
        "生产就绪（约 3–4 周）：+ Docker Backend + 审计 + fail_closed。"
    )

    # 10
    doc.add_heading("十、验收清单", level=1)
    checklist = [
        "用户 A 的 shell 无法读写用户 B 的 users/{user_id}/。",
        "文件工具绝对路径越界返回明确错误。",
        "审批通过的 rm -rf / 仅影响 sandbox_root。",
        "Docker 模式下容器无 host network、无 privileged。",
        "沙箱启动失败且 fail_closed=true 时工具不执行。",
        "MCP 工具在沙箱内运行，宿主看不到 MCP 写的越界文件。",
        "审计日志可检索：agent_id、user_id、tool、backend、duration、exit_code。",
        "approval_level=OFF 仍受 Path Jail 约束（策略可配置）。",
    ]
    for i, item in enumerate(checklist, 1):
        doc.add_paragraph(f"{i}. {item}")

    # 11
    doc.add_heading("十一、结论", level=1)
    add_table(
        doc,
        ["维度", "现状", "目标"],
        [
            ["安全模型", "策略 + 审批", "策略 + 审批 + OS 级隔离"],
            ["文件边界", "约定 workspace", "强制 sandbox_root"],
            ["执行", "宿主同进程/子进程", "SandboxRunner 统一网关"],
            ["多租户", "Agent 级", "Agent + User 子目录"],
            ["产品定位", "规划中", "Local MVP → Docker 生产"],
        ],
    )

    doc.add_paragraph()
    conclusion = doc.add_paragraph()
    conclusion.add_run("推荐路径：").bold = True
    conclusion.add_run(
        "与用户隔离方案同链路——先定 users/{user_id}/ 路径（隔离 Phase 3），"
        "再在其上挂 SandboxRunner（本方案 Phase 1–3）。"
    )

    doc.save(out_path)
    print(f"Exported: {out_path}")


if __name__ == "__main__":
    main()
