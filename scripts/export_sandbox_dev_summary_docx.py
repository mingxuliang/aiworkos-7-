# -*- coding: utf-8 -*-
"""Export today's sandbox development summary to Word (.docx)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def set_cell_shading(cell, fill: str) -> None:
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(
        qn("w:shd"),
        {qn("w:fill"): fill, qn("w:val"): "clear"},
    )
    shading.append(shd)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_fill: str = "D9E2F3",
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], header_fill)
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def main() -> None:
    out_path = Path(__file__).resolve().parent.parent / "沙箱功能开发总结.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = doc.add_heading("QwenPaw 工具沙箱功能开发总结与方案说明", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "文档类型：开发总结 / 技术方案  |  版本：v2.0  |  更新日期：2026-05-27"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph(
        "本文档汇总 QwenPaw 工具执行沙箱（Execution Sandbox）的完整开发工作。"
        "2026-05-26 完成 Path Jail、Docker Per-Call、聊天页环境切换等 MVP；"
        "2026-05-27 完成 Phase 4/5（Skill 沙箱联动、Session 容器复用、MCP Docker、状态 API）"
        "以及方案 A（Skill 只读 / 用户目录可写的读写分离 Path Jail）。"
        "文档从问题背景、总体架构、模块实现、测试验收到后续规划进行系统说明。"
    )

    # 1
    doc.add_heading("一、开发概述", level=1)
    doc.add_paragraph(
        "整体目标：在 QwenPaw / AI Work OS 中落地「工具执行沙箱」，"
        "使 Agent 在调用 read_file、write_file、execute_shell_command 等工具时，"
        "默认受限于可配置的沙箱根目录（Path Jail），并支持 Docker 按次或按 Session 复用隔离。"
        "同时在 Console 聊天页提供「沙箱环境 / 本地环境」切换，按消息粒度覆盖全局配置。"
        "2026-05-27 进一步解决 Skill 文件读取与用户写入目录分离的问题。"
    )
    add_table(
        doc,
        ["能力", "状态", "说明"],
        [
            ["Path Jail（本地路径监狱）", "已实现", "文件工具 + Tool Guard 双重拦截"],
            ["读写分离 Path Jail（方案 A）", "已实现", "skills/ 只读，users/{id}/ 可写"],
            ["Docker Per-Call Shell", "已实现", "docker run --rm 单次 shell 执行"],
            ["Session 级容器复用", "已实现", "docker exec 复用 + idle 销毁"],
            ["用户子目录 sandbox_root", "已实现", "workspaces/.../users/{user_id}/"],
            ["Skill requires_sandbox 元数据", "已实现", "Scanner 联动 + runtime 强制"],
            ["MCP Docker stdio（opt-in）", "已实现", "run_in_sandbox 时在 Session 容器内 spawn"],
            ["沙箱状态 API + Console 展示", "已实现", "Docker 健康 + Session 容器列表"],
            ["聊天页环境切换", "已实现", "execution_sandbox_enabled 按消息传递"],
            ["控制台 Security 配置", "已实现", "enabled / backend / Docker / Session / Skill 策略"],
            ["Skill 独立容器化", "不采用", "安全收益有限，改用只读 Path Jail"],
        ],
        header_fill="E2EFDA",
    )

    # 2
    doc.add_heading("二、问题背景与设计目标", level=1)

    doc.add_heading("2.1 改造前的问题", level=2)
    add_bullets(
        doc,
        [
            "Tool Guard、File Guard、Approval 属于「策略层」，工具仍在宿主 OS 用户权限下执行。",
            "read_file 使用绝对路径时可读取 workspace 外系统文件（如 C:\\Windows\\win.ini）。",
            "Shell 子进程 cwd 虽默认为 workspace，但无 OS 级强制边界。",
            "缺少与产品对标能力：聊天时可选择沙箱环境或本地环境。",
        ],
    )

    doc.add_heading("2.2 设计原则", level=2)
    add_bullets(
        doc,
        [
            "策略 + 隔离互补：保留 Tool Guard / Approval，沙箱在其之后强制执行路径边界。",
            "失败安全：Docker 不可用且 fail_closed=true 时拒绝 shell 执行。",
            "按消息可切换：聊天页切换仅影响下一条及后续消息，无需重启会话。",
            "多租户就绪：sandbox_root 支持 users/{user_id}/ 子目录。",
            "可观测：Runner 日志、Agent 环境上下文、Security 状态 API。",
        ],
    )

    doc.add_heading("2.3 非目标（首版不做）", level=2)
    add_bullets(
        doc,
        [
            "VM 级隔离、完整网络 egress 代理。",
            "Skill 脚本进程内执行改为全面容器化。",
            "Browser / Playwright 与工具沙箱合并为同一容器。",
        ],
    )

    # 3
    doc.add_heading("三、总体架构", level=1)

    doc.add_heading("3.1 三层隔离模型", level=2)
    add_table(
        doc,
        ["层级", "机制", "实现模块"],
        [
            ["L1 策略层", "Tool Guard 规则、Approval 审批", "已有 tool_guard / approval"],
            ["L2 路径层", "强制 sandbox_root，拒绝越界路径", "path_jail.py + PathJailGuardian + file_io"],
            ["L3 执行层", "Shell 在 sandbox_root 或 Docker 容器内执行", "shell.py + docker_runner.py"],
        ],
    )

    doc.add_heading("3.2 执行流水线", level=2)
    doc.add_paragraph(
        "用户消息（Console Chat）"
        " → POST /console/chat（含 execution_sandbox_enabled）"
        " → console router 写入 native_payload.meta"
        " → Runner 解析 channel_meta 并 set_sandbox_enabled_override()"
        " → load_sandbox_settings() 合并 config / env / override"
        " → resolve_sandbox_root() 计算沙箱根目录"
        " → ReactAgent 工具调用前 set_current_sandbox_root()"
        " → file_io / shell / PathJailGuardian 执行边界校验"
        " → 结果回传 Agent。"
    )

    doc.add_heading("3.3 沙箱根目录策略", level=2)
    add_code_block(
        doc,
        "workspaces/{agent_id}/                    → Agent workspace（逻辑工作区）\n"
        "workspaces/{agent_id}/skills/             → 只读根（sandbox_readonly_roots）\n"
        "workspaces/{agent_id}/users/{user_id}/    → sandbox_root（RW，工具写入边界）\n"
        "\n"
        "当 use_user_subdir=true 且 user_id 非 default 时，"
        "writable sandbox_root = workspace/users/{sanitized_user_id}/\n"
        "skills/ 及已启用 Skill 目录可通过 readonly roots 读取，但不可写入。",
    )

    doc.add_heading("3.4 Backend 方案", level=2)
    add_table(
        doc,
        ["Backend", "适用场景", "Shell 行为", "文件工具"],
        [
            ["off", "沙箱关闭", "宿主 subprocess，cwd=workspace", "无 Path Jail"],
            ["local", "开发 / 无 Docker", "宿主 subprocess，cwd=sandbox_root", "Path Jail 读写分离"],
            ["docker", "生产 / 有 Docker", "Session 容器 docker exec 或 per-call", "Path Jail + 容器内 shell"],
        ],
    )

    # 4
    doc.add_heading("四、后端核心实现", level=1)

    doc.add_heading("4.1 模块结构", level=2)
    add_table(
        doc,
        ["路径", "职责"],
        [
            ["src/qwenpaw/security/sandbox/context.py", "ContextVar：sandbox_root、enabled override、readonly roots、skill requires_sandbox、session key"],
            ["src/qwenpaw/security/sandbox/path_jail.py", "读写分离：is_path_readable/writable、resolve_tool_path_string"],
            ["src/qwenpaw/security/sandbox/resolver.py", "resolve_sandbox_root、sanitize_user_id"],
            ["src/qwenpaw/security/sandbox/settings.py", "load_sandbox_settings（config+env+override）"],
            ["src/qwenpaw/security/sandbox/docker_runner.py", "DockerSandboxRunner 按次 shell"],
            ["src/qwenpaw/security/sandbox/session_container_manager.py", "Session 容器 acquire/exec/release/idle reap"],
            ["src/qwenpaw/security/sandbox/session_container.py", "Session 容器模型与 session key 构建"],
            ["src/qwenpaw/security/sandbox/status.py", "Docker 健康 + Session 容器运行时快照"],
            ["src/qwenpaw/security/tool_guard/guardians/path_jail_guardian.py", "Tool Guard 读写模式路径拦截"],
            ["src/qwenpaw/agents/tools/file_io.py", "读写前 resolve_tool_path_string(mode=read|write)"],
            ["src/qwenpaw/agents/tools/shell.py", "local cwd / docker per-call / session docker exec"],
            ["src/qwenpaw/app/mcp/docker_stdio_client.py", "MCP stdio 在 Session 容器内 docker exec -i"],
            ["src/qwenpaw/agents/skills_manager.py", "SkillRequirements.requires_sandbox 元数据"],
            ["src/qwenpaw/security/skill_scanner/", "recommend_sandbox、auto_tag_risky_skills"],
            ["src/qwenpaw/agents/react_agent.py", "每次工具轮次设置 sandbox_root"],
            ["src/qwenpaw/app/runner/runner.py", "override、Skill 强制、Session 容器、readonly roots 生命周期"],
            ["src/qwenpaw/app/_app.py", "Session 容器 idle reaper 后台任务"],
            ["src/qwenpaw/app/routers/console.py", "AgentRequest 读取 execution_sandbox_enabled"],
        ],
    )

    doc.add_heading("4.2 启用判定优先级", level=2)
    doc.add_paragraph(
        "is_sandbox_enabled() 与 load_sandbox_settings() 按以下顺序决定当前请求是否启用沙箱："
    )
    add_bullets(
        doc,
        [
            "最高：ContextVar current_sandbox_enabled_override（来自聊天请求 execution_sandbox_enabled）。",
            "override=False 时直接关闭沙箱；override=True 时强制开启（backend 仍读 config/env）。",
            "其次：环境变量 QWENPAW_EXECUTION_SANDBOX_ENABLED。",
            "最后：config.json → security.execution_sandbox.enabled 且 backend != off。",
        ],
    )

    doc.add_heading("4.3 Path Jail 关键逻辑", level=2)
    add_bullets(
        doc,
        [
            "读写分离（方案 A）：writable root = users/{user_id}/；readonly roots = skills/ + 已启用 Skill 目录。",
            "is_path_readable：writable root 或任一 readonly root 内路径允许读取。",
            "is_path_writable：仅 writable root 内路径允许写入。",
            "resolve_tool_path_string(mode=read|write)：相对路径 skills/... 基于 workspace 解析。",
            "Windows 使用大小写不敏感前缀比较，避免 C:\\ 与 c:\\ 绕过。",
            "越界抛出 SandboxBoundaryError，file_io 转为 ToolResponse 错误文本。",
            "PathJailGuardian：write 工具用写策略，read/shell 用读策略。",
        ],
    )

    doc.add_heading("4.4 Skill 沙箱联动（Phase 4）", level=2)
    add_bullets(
        doc,
        [
            "Skill 元数据新增 requires_sandbox（SkillRequirements + manifest requires.sandbox）。",
            "Skill Scanner 扫描结果 recommend_sandbox，auto_tag_risky_skills 可自动写入 manifest。",
            "skill_sandbox_enforcement：off / warn / strict；strict 模式下未开沙箱时阻止启用高风险 Skill。",
            "Runner 检测 /skill 指令或 Skill 调用，set_current_skill_requires_sandbox(True) 强制沙箱。",
            "is_sandbox_enabled() 最高优先级之一：current_skill_requires_sandbox=True 时强制开启。",
            "结论：requires_sandbox 强制进程内 Agent 的工具沙箱，而非 Skill 独立容器化。",
        ],
    )

    doc.add_heading("4.5 Session 容器复用（Phase 5）", level=2)
    add_bullets(
        doc,
        [
            "SessionContainerManager：按 agent_id:user_id:session_id 键复用 Docker 容器。",
            "Runner 请求开始时 acquire，结束时 release；shell 走 docker exec 而非每次 run --rm。",
            "配置：session_container_enabled、session_idle_seconds（默认 900s）、session_max_containers（默认 32）。",
            "_app.py 后台 idle reaper 定期销毁超时容器。",
            "Status API 返回 active_count、containers 列表、idle_seconds。",
        ],
    )

    doc.add_heading("4.6 MCP Docker stdio（Phase 4）", level=2)
    add_bullets(
        doc,
        [
            "DockerStdIOStatefulClient：MCP Server 在 Session 容器内通过 docker exec -i 启动。",
            "MCP 配置 run_in_sandbox=true 时 opt-in 启用（非默认）。",
            "依赖 Session 容器已 acquire，与 shell 共用同一 sandbox_root 挂载。",
        ],
    )

    doc.add_heading("4.7 Docker Per-Call（Plan B）", level=2)
    add_bullets(
        doc,
        [
            "镜像：qwenpaw-sandbox:latest（deploy/Dockerfile.sandbox，基于 debian:bookworm-slim）。",
            "命令：docker run --rm --network none --memory 512m --cpus 1 --pids-limit 64。",
            "挂载：sandbox_root → /work:rw；若 workspace ≠ sandbox_root 则 workspace → /ro:ro。",
            "输出前缀：[sandbox:docker X.XXs] 便于日志与测试识别。",
            "fail_closed=true 且 Docker 不可用时返回错误，不静默降级到宿主 shell。",
        ],
    )
    doc.add_paragraph("构建镜像：")
    add_code_block(doc, ".\\scripts\\build_sandbox_image.ps1")

    doc.add_heading("4.8 Agent 环境上下文", level=2)
    doc.add_paragraph(
        "build_env_context()（runner/utils.py）向 Agent 注入当前执行环境说明，避免模型误判："
    )
    add_bullets(
        doc,
        [
            "沙箱模式：Execution environment: Sandbox (backend=local|docker)、Sandbox root 路径。",
            "本地模式：Execution environment: Local (sandbox disabled)。",
            "明确说明：Host OS 信息是后端宿主机，非独立远程 VM。",
        ],
    )

    doc.add_heading("4.9 关键 Bug 修复", level=2)
    doc.add_paragraph(
        "问题：聊天页显示沙箱模式，但 Agent 仍可读取 C:\\Windows\\win.ini。"
    )
    doc.add_paragraph(
        "根因：FastAPI 将 POST body 解析为 AgentRequest（Pydantic），"
        "_extract_session_and_payload() 原先仅从 dict 读取 execution_sandbox_enabled，"
        "导致 override 未进入 channel_meta，Path Jail 未生效。"
    )
    doc.add_paragraph(
        "修复：新增 _read_execution_sandbox_enabled()，同时支持 AgentRequest 与 dict；"
        "Runner 从 channel_meta 读取并 set_sandbox_enabled_override()。"
        "单元测试：tests/unit/routers/test_console_placeholder.py。"
    )

    # 5
    doc.add_heading("五、配置与 API", level=1)

    doc.add_heading("5.1 config.json 配置项", level=2)
    add_code_block(
        doc,
        '{\n'
        '  "security": {\n'
        '    "execution_sandbox": {\n'
        '      "enabled": true,\n'
        '      "backend": "local",\n'
        '      "use_user_subdir": true,\n'
        '      "fail_closed": true,\n'
        '      "fallback_backend": "local",\n'
        '      "docker_image": "qwenpaw-sandbox:latest",\n'
        '      "docker_network": "none",\n'
        '      "docker_memory": "512m",\n'
        '      "docker_cpus": "1",\n'
        '      "docker_pids_limit": 64,\n'
        '      "docker_timeout_seconds": 120,\n'
        '      "skill_sandbox_enforcement": "warn",\n'
        '      "auto_tag_risky_skills": true,\n'
        '      "session_container_enabled": false,\n'
        '      "session_idle_seconds": 900,\n'
        '      "session_max_containers": 32,\n'
        '      "sandbox_readonly_roots": ["skills"],\n'
        '      "allow_enabled_skill_dirs": true\n'
        '    }\n'
        '  }\n'
        '}',
    )

    doc.add_heading("5.2 环境变量", level=2)
    add_table(
        doc,
        ["变量", "作用"],
        [
            ["QWENPAW_EXECUTION_SANDBOX_ENABLED", "true/false，覆盖 config.enabled"],
            ["QWENPAW_EXECUTION_SANDBOX_BACKEND", "local / docker / off"],
        ],
    )

    doc.add_heading("5.3 HTTP API", level=2)
    add_table(
        doc,
        ["接口", "说明"],
        [
            ["GET /api/config/security/execution-sandbox", "读取沙箱配置"],
            ["PUT /api/config/security/execution-sandbox", "更新沙箱配置"],
            ["GET /api/config/security/execution-sandbox/status", "Docker 可用性、镜像、Session 容器运行时状态"],
            ["POST /console/chat", "请求体字段 execution_sandbox_enabled: boolean，按消息覆盖"],
        ],
    )

    # 6
    doc.add_heading("六、前端实现（Console 聊天页）", level=1)

    doc.add_heading("6.1 交互设计", level=2)
    add_bullets(
        doc,
        [
            "输入栏左侧 prefix 区域放置环境选择器，与附件、麦克风图标风格一致（Spark 图标）。",
            "沙箱环境：SparkComputerLine；本地环境：SparkLocalFileLine。",
            "点击下拉选择「沙箱环境」或「本地环境」，选择持久化到 localStorage。",
            "默认：沙箱环境（qwenpaw.chat.execution_sandbox_enabled = true）。",
            "会话中可随时切换，下一条消息起生效（按消息粒度，无需灰显或锁定）。",
        ],
    )

    doc.add_heading("6.2 关键文件", level=2)
    add_table(
        doc,
        ["路径", "说明"],
        [
            ["console/src/hooks/useExecutionEnvironment.ts", "mode 状态 + localStorage"],
            [
                "console/src/pages/Chat/components/ExecutionEnvironmentSelector/",
                "Dropdown + IconButton 组件",
            ],
            [
                "console/src/pages/Chat/components/ChatSenderPrefixActions/",
                "输入栏 prefix 动作区",
            ],
            ["console/src/pages/Chat/index.tsx", "customFetch 传递 execution_sandbox_enabled"],
            ["console/src/pages/Settings/Security/components/ExecutionSandboxSection.tsx", "沙箱配置 + Session 容器 + Skill 策略 + 状态展示"],
            ["console/src/locales/zh.json / en.json", "chat.executionEnvironment.* 文案"],
        ],
    )

    doc.add_heading("6.3 请求传参", level=2)
    add_code_block(
        doc,
        "// 每条聊天消息 POST body 附带\n"
        "{\n"
        '  "session_id": "...",\n'
        '  "user_id": "...",\n'
        '  "execution_sandbox_enabled": true,  // 或 false\n'
        '  "input": [...]\n'
        "}",
    )

    # 7
    doc.add_heading("七、测试与验收", level=1)

    doc.add_heading("7.1 自动化脚本", level=2)
    add_table(
        doc,
        ["脚本", "验证内容"],
        [
            ["scripts/sandbox_smoke_test.py", "PathJailGuardian 冒烟"],
            ["scripts/sandbox_integration_check.py", "ToolGuardEngine 集成"],
            ["scripts/sandbox_e2e_agent_test.py", "read/write 工具 E2E"],
            ["scripts/verify_execution_environment_toggle.py", "聊天切换 + AgentRequest + health"],
            ["scripts/sandbox_console_test.py", "config 持久化 + local/docker 场景"],
            ["scripts/sandbox_docker_check.py", "Docker shell 集成"],
        ],
    )
    doc.add_paragraph("详细说明见同目录文档：沙箱测试脚本.docx")

    doc.add_heading("7.2 单元测试", level=2)
    add_code_block(
        doc,
        "pytest tests/unit/security/test_path_jail.py "
        "tests/unit/security/test_docker_runner.py "
        "tests/unit/security/test_sandbox_status.py "
        "tests/unit/security/test_skill_sandbox.py "
        "tests/unit/security/test_session_container_manager.py "
        "tests/unit/routers/test_console_placeholder.py -q",
    )
    add_bullets(
        doc,
        [
            "test_path_jail.py（19 项）：sandbox_root 解析、穿越拦截、用户隔离、Skill 只读/不可写、file_io、Guardian",
            "test_skill_sandbox.py（5 项）：requires_sandbox 元数据、Scanner recommend_sandbox",
            "test_docker_runner.py：run 命令组装、挂载、超时",
            "test_sandbox_status.py：Docker 健康 + Session 容器状态",
            "test_session_container_manager.py：acquire 复用、idle reap",
            "console：AgentRequest / dict 的 execution_sandbox_enabled 提取",
        ],
    )

    doc.add_heading("7.3 验收标准（已通过）", level=2)
    add_bullets(
        doc,
        [
            "沙箱模式：read_file(\"C:/Windows/win.ini\") 返回 outside sandbox 错误。",
            "本地模式：同上路径可正常读取（override=false）。",
            "沙箱内 write_file + read_file 正常；../ 路径穿越写被拦截。",
            "用户 A 无法读取用户 B 的 users/{id}/ 私有文件。",
            "Skill 文件 skills/.../SKILL.md 可从 users/{id}/ 沙箱上下文直接 read_file。",
            "write_file 到 skills/ 目录被拦截。",
            "Docker 模式：shell 输出含 [sandbox:docker 与 /work 工作目录。",
            "聊天页切换后下一条消息行为与选择一致。",
            "方案 A 无需 Docker，Local backend 即可验证 Skill 只读访问。",
        ],
    )

    # 8
    doc.add_heading("八、交付清单", level=1)
    add_table(
        doc,
        ["类别", "交付物"],
        [
            ["后端 MVP（05-26）", "sandbox 模块、Path Jail、Docker Runner、Runner/Console 传参、env_context"],
            ["后端 Phase 4/5（05-27）", "Skill requires_sandbox、SessionContainerManager、MCP Docker、Status API、idle reaper"],
            ["后端方案 A（05-27）", "读写分离 Path Jail、readonly roots、runner 生命周期注入 Skill 目录"],
            ["前端", "ExecutionEnvironmentSelector、ExecutionSandboxSection 扩展、Session 状态展示"],
            ["配置", "ExecutionSandboxConfig 全字段、Security API、控制台 Security 页"],
            ["测试", "6+ 脚本 + 26+ 单元测试（path_jail/skill/session/status）"],
            ["文档", "工具沙箱开发方案.docx、沙箱测试脚本.docx、本文档"],
            ["镜像", "deploy/Dockerfile.sandbox、build_sandbox_image.ps1"],
        ],
        header_fill="FFF2CC",
    )

    # 9
    doc.add_heading("九、与 Tool Guard 的关系", level=1)
    add_table(
        doc,
        ["场景", "Tool Guard", "沙箱 Path Jail"],
        [
            ["读 /etc/passwd 或 win.ini", "可能未命中规则", "硬拒绝（沙箱开启时）"],
            ["rm -rf 危险命令", "规则命中 → 审批", "即使批准也只能操作 sandbox_root 内"],
            ["../ 路径穿越", "部分规则 WARN", "resolve_path_in_jail 硬拒绝"],
            ["本地模式（override=false）", "仍生效", "Path Jail 关闭，恢复宿主权限"],
        ],
    )
    doc.add_paragraph(
        "结论：Tool Guard 是「门卫」，沙箱是「牢房」。两者叠加，不可互相替代。"
    )

    # 10
    doc.add_heading("十、已知限制", level=1)
    add_bullets(
        doc,
        [
            "local backend 仍共享 QwenPaw 进程 OS 用户，防误操作不防内核级攻击。",
            "Agent / ReAct 推理仍在宿主进程内，requires_sandbox 仅强制工具层沙箱。",
            "Skill 脚本未独立容器化（经评估安全收益有限，采用只读 Path Jail 替代）。",
            "MCP run_in_sandbox 为 opt-in，默认仍在宿主 spawn。",
            "Session 容器与 Docker Per-Call 需 Docker Desktop / WSL2（Windows 开发环境）。",
            "Browser / Playwright 与工具沙箱未合并为同一容器。",
            "用户隔离 Phase 0 路径策略与沙箱共用 resolver，多租户完整验收待用户目录全面落地。",
        ],
    )

    # 11
    doc.add_heading("十一、后续规划", level=1)
    add_table(
        doc,
        ["阶段", "内容", "优先级"],
        [
            ["Phase 6", "结构化沙箱审计日志、检索 API", "P2"],
            ["Phase 6", "OpenShell / 更强网络 egress 策略", "P3"],
            ["产品", "切换环境时 Toast「下一条消息起生效」", "P3"],
            ["产品", "/environment 命令查询当前模式", "P3"],
            ["产品", "Console 配置 sandbox_readonly_roots 可视化编辑", "P3"],
            ["运维", "Session 容器监控告警、镜像版本管理", "P3"],
        ],
    )

    # 12
    doc.add_heading("十二、2026-05-27 增量开发详述", level=1)

    doc.add_heading("12.1 问题与决策", level=2)
    doc.add_paragraph(
        "启用 per-user 沙箱（users/81/）后，Skill 文件位于 workspace/skills/，"
        "与 sandbox_root 为兄弟目录，导致 read_file 读取 Skill 脚本/references 被 Path Jail 拦截，"
        "Agent 被迫用 shell cp 绕行。"
    )
    doc.add_paragraph(
        "曾评估 Skill 独立容器化方案，结论为：Agent 仍在宿主 ReAct 循环内，"
        "单独容器化 Skill 无法隔离 LLM 推理与工具调用链，安全收益有限、复杂度高。"
        "最终采用方案 A：读写分离 Path Jail。"
    )

    doc.add_heading("12.2 方案 A 实现要点", level=2)
    add_table(
        doc,
        ["组件", "改动"],
        [
            ["config.py", "sandbox_readonly_roots、allow_enabled_skill_dirs"],
            ["path_jail.py", "is_path_readable/writable、resolve_tool_path_string(mode)"],
            ["context.py", "current_readonly_roots ContextVar"],
            ["file_io.py", "_resolve_read_path / _resolve_write_path 分离"],
            ["path_jail_guardian.py", "read/write 模式分别校验"],
            ["runner.py", "Agent 初始化后注入 enabled skill dirs，finally 清理"],
        ],
        header_fill="E2EFDA",
    )

    doc.add_heading("12.3 Phase 4/5 PR 拆分", level=2)
    add_table(
        doc,
        ["PR", "内容", "关键文件"],
        [
            ["PR-1", "Skill requires_sandbox + Scanner 联动 + runtime 强制", "skills_manager.py、skill_scanner、runner.py"],
            ["PR-2", "SessionContainerManager + MCP docker stdio", "session_container_manager.py、docker_stdio_client.py"],
            ["PR-3", "Shell session docker exec 复用 + idle reaper", "shell.py、_app.py、runner.py"],
            ["PR-4", "Status API + Console 展示", "status.py、ExecutionSandboxSection.tsx"],
        ],
        header_fill="D9E2F3",
    )

    doc.add_heading("12.4 目录布局（方案 A 生效后）", level=2)
    add_code_block(
        doc,
        "workspaces/default/\n"
        "  skills/              ← 只读（sandbox_readonly_roots）\n"
        "    demo-skill/\n"
        "      SKILL.md         ← read_file 允许\n"
        "      scripts/         ← read_file 允许\n"
        "  users/\n"
        "    81/                ← 可写 sandbox_root\n"
        "      output.txt       ← write_file 允许\n"
        "\n"
        "Agent 在 users/81/ 沙箱内运行，可直接读取 skills/ 下 Skill 资源，"
        "但不能修改 skills/ 或访问其他用户目录。",
    )

    doc.add_heading("12.5 手动验证建议", level=2)
    add_bullets(
        doc,
        [
            "开启 Local 沙箱（无需 Docker）：设置 → 安全 → 执行沙箱 → enabled + backend=local。",
            "调用读取 skills/ 下 scripts 或 references 的 Skill，确认 read_file 直接成功。",
            "尝试 write_file 到 skills/，应返回 outside writable sandbox 错误。",
            "Docker 相关（Session 容器、MCP Docker）需 Docker Desktop 运行。",
        ],
    )

    # 13
    doc.add_heading("十三、附录：数据流时序", level=1)
    doc.add_paragraph(
        "Console 用户选择「沙箱环境」并发送消息 → "
        "Chat index.tsx customFetch 设置 execution_sandbox_enabled=true → "
        "console.py _extract_session_and_payload 写入 meta → "
        "Runner handle_agent_query 调用 set_sandbox_enabled_override(True) → "
        "load_sandbox_settings() 返回 enabled=true → "
        "resolve_sandbox_root(workspace, user_id) → "
        "Runner 注入 readonly roots（skills/ + enabled skill dirs）→ "
        "ReactAgent 工具调用 set_current_sandbox_root → "
        "read_file 调用 resolve_tool_path_string(mode=read) → "
        "skills/ 路径可读、users/{id}/ 外路径拒绝 → "
        "write_file 调用 mode=write，skills/ 写入拒绝。"
    )

    doc.add_heading("附录：相关文档", level=1)
    add_bullets(
        doc,
        [
            "工具沙箱开发方案.docx — 分阶段规划与 Backend 对比（规划文档）",
            "沙箱测试脚本.docx — 测试脚本使用手册",
            "scripts/export_sandbox_plan_docx.py — 规划文档导出脚本",
            "scripts/export_sandbox_test_scripts_docx.py — 测试手册导出脚本",
            "scripts/export_sandbox_dev_summary_docx.py — 本文档导出脚本",
        ],
    )

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
