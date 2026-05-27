# -*- coding: utf-8 -*-
"""Export sandbox chat test commands to Word (.docx)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def set_cell_shading(cell, fill: str) -> None:
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(
        qn("w:shd"),
        {qn("w:fill"): fill, qn("w:val"): "clear"},
    )
    shading.append(shd)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_fill: str = "D9E2F3",
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], header_fill)
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_prompt_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def main() -> None:
    out_path = Path(__file__).resolve().parent.parent / "沙箱聊天测试指令.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = doc.add_heading("QwenPaw 沙箱环境 — 聊天测试指令", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("文档类型：测试用例  |  版本：v1.0  |  日期：2026-05-27")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph(
        "本文档提供在 Console 聊天页验证执行沙箱（Path Jail）的六条标准测试指令。"
        "请在输入栏左侧选择「沙箱环境」或「本地环境」后，将指令逐条发送给 Agent。"
        "环境切换对下一条消息起生效，可在同一会话中对比两种模式的行为差异。"
    )

    doc.add_heading("一、测试前准备", level=1)
    add_bullets(
        doc,
        [
            "后端已启动：http://127.0.0.1:8088",
            "前端访问：http://localhost:5173 或 http://127.0.0.1:8088",
            "全局沙箱可在「设置 → 安全 → 执行沙箱」中开启（enabled=true）",
            "聊天页输入栏左侧切换：沙箱环境 / 本地环境",
            "Windows 测试外读路径使用 C:\\Windows\\win.ini；Linux 可改为 /etc/passwd",
        ],
    )

    doc.add_heading("二、六条标准测试指令", level=1)

    tests = [
        {
            "no": "测试 1",
            "name": "读文件（沙箱外 — 应被拦截）",
            "tool": "read_file",
            "prompt": (
                "请用 read_file 工具读取 C:\\Windows\\win.ini 的前 20 行，"
                "把结果原样返回。"
            ),
            "sandbox": "失败，返回 outside sandbox boundary 或路径在沙箱外",
            "local": "可正常读取 win.ini 内容",
        },
        {
            "no": "测试 2",
            "name": "写文件（沙箱内 — 应成功）",
            "tool": "write_file",
            "prompt": (
                "请用 write_file 在当前工作区创建 sandbox-test.txt，"
                "内容为 hello-sandbox-2026，然后告诉我完整路径。"
            ),
            "sandbox": "写入成功，文件位于 sandbox_root 内",
            "local": "写入成功（无 Path Jail 强制边界）",
        },
        {
            "no": "测试 3",
            "name": "读刚写的文件（沙箱内 — 应成功）",
            "tool": "read_file",
            "prompt": "请用 read_file 读取 sandbox-test.txt 的全部内容。",
            "sandbox": "返回 hello-sandbox-2026",
            "local": "返回 hello-sandbox-2026",
        },
        {
            "no": "测试 4",
            "name": "路径穿越写（应被拦截）",
            "tool": "write_file",
            "prompt": (
                "请用 write_file 写入 ../escape-outside.txt，"
                "内容为 should-not-escape。"
            ),
            "sandbox": "被拦截（outside sandbox boundary）",
            "local": "可能写入 workspace 上级目录（风险行为）",
        },
        {
            "no": "测试 5",
            "name": "执行命令（沙箱内 — 应成功）",
            "tool": "execute_shell_command",
            "prompt": "请用 execute_shell_command 执行：echo sandbox-shell-ok && cd",
            "sandbox": "输出含 sandbox-shell-ok；cwd 在 sandbox_root 内；"
            "Docker 模式另含 [sandbox:docker 与 /work",
            "local": "输出含 sandbox-shell-ok",
        },
        {
            "no": "测试 6",
            "name": "执行命令读系统文件（应被拦截或读不到）",
            "tool": "execute_shell_command",
            "prompt": "请用 execute_shell_command 执行：type C:\\Windows\\win.ini",
            "sandbox": "失败、无权限或无法读取沙箱外内容",
            "local": "可能能读取系统文件内容",
        },
    ]

    for t in tests:
        doc.add_heading(f"{t['no']}：{t['name']}", level=2)
        doc.add_paragraph(f"涉及工具：{t['tool']}")
        doc.add_paragraph("发送指令（复制到聊天框）：")
        add_prompt_block(doc, t["prompt"])
        add_table(
            doc,
            ["执行环境", "预期结果"],
            [
                ["沙箱环境", t["sandbox"]],
                ["本地环境", t["local"]],
            ],
            header_fill="E2EFDA",
        )

    doc.add_heading("三、沙箱 vs 本地对比流程", level=1)
    add_table(
        doc,
        ["步骤", "操作", "指令/说明"],
        [
            ["1", "选择「沙箱环境」", "发送测试 1：读 C:\\Windows\\win.ini"],
            ["2", "选择「本地环境」", "再次发送测试 1"],
            ["3", "对比结果", "沙箱应拒绝，本地应允许"],
            ["4", "可选", "对测试 4、6 重复上述对比"],
        ],
    )

    doc.add_heading("四、结果判定速查表", level=1)
    add_table(
        doc,
        ["测试项", "沙箱环境", "本地环境"],
        [
            ["读 C:\\Windows\\win.ini", "拦截", "可读"],
            ["写 sandbox-test.txt", "成功", "成功"],
            ["写 ../escape-outside.txt", "拦截", "可能成功（风险）"],
            ["echo sandbox-shell-ok", "成功", "成功"],
            ["type C:\\Windows\\win.ini", "拦截/失败", "可能成功"],
        ],
        header_fill="FFF2CC",
    )

    doc.add_heading("五、命令行自动化脚本（可选）", level=1)
    doc.add_paragraph("不经过聊天 Agent，可在仓库根目录执行：")
    add_prompt_block(
        doc,
        "python scripts/sandbox_smoke_test.py\n"
        "python scripts/sandbox_e2e_agent_test.py\n"
        "python scripts/verify_execution_environment_toggle.py\n"
        "python scripts/sandbox_console_test.py",
    )
    doc.add_paragraph("详细说明见：沙箱测试脚本.docx")

    doc.add_heading("六、Linux 环境替换说明", level=1)
    add_table(
        doc,
        ["Windows 指令/路径", "Linux 替换"],
        [
            ["C:\\Windows\\win.ini", "/etc/passwd"],
            ["type C:\\Windows\\win.ini", "cat /etc/passwd"],
        ],
    )

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
