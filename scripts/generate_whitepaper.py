#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate ai-work-os product marketing white paper (Word, v3.1)."""

from __future__ import annotations

import urllib.request
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BRAND = "ai-work-os"
ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
ASSETS = DOCS / "whitepaper_assets"
OUTPUT = DOCS / f"{BRAND}_产品白皮书_v3.1_行业案例.docx"
OUTPUT_EN = DOCS / f"{BRAND}_whitepaper_v3.1_cases.docx"

# Official console screenshots (from product docs)
IMAGES = {
    "console_chat": (
        "https://img.alicdn.com/imgextra/i1/O1CN01ikrU3k1TRdNESHtzV_!!6000000002379-2-tps-3822-2070.png",
        "控制台对话界面：企业成员通过 Web 与 Agent 协作",
    ),
    "security": (
        "https://img.alicdn.com/imgextra/i4/O1CN01APDb151R02HB3zHet_!!6000000002048-2-tps-3822-2070.png",
        "安全中心：工具守卫、文件防护、技能扫描、执行沙箱统一配置",
    ),
    "skills": (
        "https://img.alicdn.com/imgextra/i1/O1CN015rwm8V24Aoc48bGTp_!!6000000007351-2-tps-3822-2070.png",
        "技能管理：Skill 上架前扫描，运行时 Policy 引擎实时评估",
    ),
    "agents": (
        "https://img.alicdn.com/imgextra/i3/O1CN01mG2dkX1GHTaJLMjWo_!!6000000000597-2-tps-3822-2070.png",
        "多智能体管理：按岗位/业务线部署独立 Agent 工作区",
    ),
    "channels": (
        "https://img.alicdn.com/imgextra/i4/O1CN01i17DLe2559KSAMwpA_!!6000000007474-2-tps-3822-2070.png",
        "频道接入：钉钉、飞书、企微等 IM 一键接入，业务触达零改造",
    ),
    "cron": (
        "https://img.alicdn.com/imgextra/i1/O1CN01UqMEHK1g2WfLs2wUQ_!!6000000004084-2-tps-3822-2070.png",
        "定时任务：7×24 自动化执行，释放人力做高价值判断",
    ),
    "files": (
        "https://img.alicdn.com/imgextra/i4/O1CN01Buqn3J1YJSNgrz9eY_!!6000000003038-2-tps-3822-2070.png",
        "工作区文件：AI 原生文件系统，操作可追溯、可备份、可回退",
    ),
    "team": (
        "https://img.alicdn.com/imgextra/i2/O1CN01I97HXk27XCpGHn6KL_!!6000000007806-2-tps-3442-1788.png",
        "多 Agent 协作：一人团队（OPT）模式下的人机分工与任务编排",
    ),
    "models": (
        "https://img.alicdn.com/imgextra/i2/O1CN0122wKIo2AIUx3HZkdG_!!6000000008180-2-tps-3822-2070.png",
        "模型管理：云端大模型 + 本地小模型协同，成本与隐私兼顾",
    ),
}


def set_run_font(run, name: str = "微软雅黑", size: int = 11, bold: bool = False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def download_images() -> dict[str, Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)
    paths: dict[str, Path] = {}
    for key, (url, _) in IMAGES.items():
        ext = ".png" if url.endswith(".png") else ".jpg"
        dest = ASSETS / f"{key}{ext}"
        if dest.exists() and dest.stat().st_size > 100_000:
            paths[key] = dest
            continue
        for attempt in range(3):
            try:
                req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
                with urllib.request.urlopen(req, timeout=120) as resp:
                    data = resp.read()
                if len(data) < 50_000:
                    raise OSError(f"response too small: {len(data)} bytes")
                dest.write_bytes(data)
                paths[key] = dest
                print(f"  OK {key} ({len(data)} bytes)")
                break
            except Exception as exc:
                print(f"  retry {key} attempt {attempt + 1}: {exc}")
        else:
            print(f"Warning: failed to download {key} after retries")
    return paths


def add_title_page(doc: Document):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(BRAND)
    set_run_font(run, size=32, bold=True, color=(0, 51, 102))

    for subtitle in [
        "企业级智能体 OPT 平台",
        "产品白皮书",
        "安全原生架构 · 一人团队 · 生产力跃升",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        set_run_font(run, size=16 if "白皮书" not in subtitle else 22, bold="白皮书" in subtitle)

    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"版本 3.1  |  {date.today().strftime('%Y年%m月')}")
    set_run_font(run, size=11, color=(120, 120, 120))
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size={1: 18, 2: 14, 3: 12}.get(level, 11), bold=True, color=(0, 51, 102))


def add_para(doc: Document, text: str, bold: bool = False, indent: bool = False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True, size=10)
    for row_data in rows:
        row = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row[i].text = cell_text
            for p in row[i].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10)
    doc.add_paragraph()


def add_image(doc: Document, img_path: Path | None, caption: str):
    if img_path and img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(6.2))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        set_run_font(r, size=9, color=(100, 100, 100))
        doc.add_paragraph()
    else:
        add_para(doc, f"[图示占位] {caption}", indent=True)


def add_case_block(
    doc: Document,
    title: str,
    meta: list[tuple[str, str]],
    background: str,
    assistants: list[str],
    value: list[str],
    platform: list[str],
):
    """Structured customer case study block."""
    add_heading(doc, title, 2)
    add_table(doc, ["项目要素", "说明"], meta)
    add_para(doc, "业务背景", bold=True)
    add_para(doc, background, indent=True)
    add_para(doc, "部署助手 / 解决方案", bold=True)
    for item in assistants:
        add_bullet(doc, item)
    add_para(doc, "客户价值", bold=True)
    for item in value:
        add_bullet(doc, item)
    add_para(doc, f"启用的 {BRAND} 能力", bold=True)
    for item in platform:
        add_bullet(doc, item)
    doc.add_paragraph()


def build_document(img_paths: dict[str, Path]) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    add_title_page(doc)

    add_heading(doc, "目录", 1)
    for item in [
        "一、核心观点：不是「有安全功能」，而是「架构在安全体系之上」",
        "二、产品创新点",
        "三、OPT 平台：One Person Team 一人团队",
        "四、六层安全体系：企业 AI 守门员",
        "五、产品整体架构",
        "六、产品功能详解（图文并茂）",
        "七、相对 OpenClaw / Hermes 的企业级优势",
        "八、行业标杆案例",
        "九、企业效益：生产力、成本、合规",
        "十、适用人群与部署建议",
        "十一、总结",
    ]:
        add_para(doc, item)
    doc.add_page_break()

    # ── 一、核心观点 ──
    add_heading(doc, "一、核心观点：不是「有安全功能」，而是「架构在安全体系之上」", 1)
    add_para(
        doc,
        "在「养虾潮」（OpenClaw 类个人 Agent）迅速升温之后，Agent 安全性成为企业用户最担心的问题。"
        "工信部曾指出，部分 OpenClaw 实例存在较高安全风险；业界也警示：若将不受约束的 Agent "
        "「全部放出来」，可能危害整个数字生态。",
        indent=True,
    )
    add_para(
        doc,
        f"与 OpenClaw、Hermes 等 To-C 个人助手「先能力、后补安全」的路径不同，"
        f"{BRAND} 从第一天起就将安全作为操作系统内核——"
        "不是给 Agent 贴一层安全补丁，而是让每一次推理、每一次工具调用、"
        "每一次文件读写，都发生在可定义、可阻断、可审计的策略边界之内。",
        indent=True,
    )
    add_para(
        doc,
        f"这正是 {BRAND} 作为企业级 OPT 平台与 To-C 个人助手的根本分野："
        "个人助手回答「我能帮你做什么」；企业平台回答「在什么边界内、"
        "由谁授权、对什么数据、可如何追溯」。",
        bold=True,
        indent=True,
    )

    # ── 二、创新点 ──
    add_heading(doc, "二、产品创新点", 1)
    innovations = [
        (
            "安全原生架构（Security-Native OS）",
            "安全不是插件，而是 Agent 运行的前置条件。"
            "Tool Guard、File Guard、Skill Scanner、Execution Sandbox "
            "构成统一 Policy 引擎，贯穿接入—编排—执行全链路。",
        ),
        (
            "OPT 人机协作岗位平台",
            "OPT = One Person Team（一人团队）。"
            "为销售、运营、人力、研发、财务等岗位预置「人做判断、AI 做执行」"
            "的协同流程，一人即一支队伍。",
        ),
        (
            "六层纵深安全体系",
            "双层规则 + 统一身份 + 容器沙箱 + Skill 供应链治理 + 全链路审计 + "
            "AI 原生文件系统，对标企业等保与内审要求。",
        ),
        (
            "多 Agent 工作区隔离",
            "每个岗位/业务线独立 Agent、独立记忆、独立技能池，"
            "支持热重载与跨 Agent 协作，避免上下文污染与数据串扰。",
        ),
        (
            "Skill 能力市场与供应链安全",
            "Skill 上架前自动扫描，运行时 Policy 逐步评估；"
            "兼容开源 Skill 生态，但以企业 Policy 为最终裁决。",
        ),
        (
            "自主研发、私有化交付",
            "Python 技术栈，支持 pip/Docker/离线本地模型；"
            "数据、密钥、审计日志均可留在客户内网，满足金融/政务/制造等行业交付。",
        ),
    ]
    for title, desc in innovations:
        add_heading(doc, title, 3)
        add_para(doc, desc, indent=True)

    # ── 三、OPT ──
    add_heading(doc, "三、OPT 平台：One Person Team（一人团队）", 1)
    add_para(
        doc,
        "OPT 并非 Orchestration·Policy·Trust 的技术缩写，而是 "
        "One Person Team——让一名员工借助 AI Agent 获得一支团队的产出能力。"
        f"{BRAND} 将 OPT 从技术概念落地为可配置的岗位协同平台："
        "每个组织节点可定义 AI 赋能深度、人员规模、人机分工步骤与专属 Agent。",
        indent=True,
    )
    add_table(
        doc,
        ["OPT 岗位域", "人工员工专注", "AI Agent 承担", "典型效益"],
        [
            ["销售", "客户关系、商务谈判、策略定价", "线索整理、竞品分析、方案初稿、跟进提醒", "销售人效 +30%~50%"],
            ["运营", "策略制定、异常仲裁、品牌调性", "数据采集、报表生成、活动排期、A/B 分析", "运营事务自动化 60%+"],
            ["人力 HR", "组织决策、员工关怀、合规判断", "简历筛选、制度问答、入离职流程、考勤汇总", "事务性 HR 工时 -40%"],
            ["研发", "架构设计、Code Review、技术决策", "代码生成、单测、文档、Issue 分拣", "交付周期缩短 20%+"],
            ["财务/财税", "合规审批、风险判断", "票据识别、对账、报表、政策检索", "关账周期压缩"],
            ["法务", "诉讼策略、合同谈判", "合同审查、法规检索、风险条款标注", "审查效率 3~5 倍"],
            ["设计", "创意方向、品牌终审", "素材生成、尺寸适配、版本管理", "出图迭代加速"],
            ["客服", "复杂投诉、舆情危机", "FAQ 应答、工单分类、多语言回复", "一线客服成本降低"],
            ["IT 运维", "变更审批、故障定责", "日志分析、脚本执行（沙箱内）、巡检", "MTTR 显著下降"],
            ["管理/CEO", "战略决策、资源仲裁", "经营简报、风险模拟、跨部门协调", "决策信息实时化"],
        ],
    )
    add_image(doc, img_paths.get("team"), IMAGES["team"][1])
    add_para(
        doc,
        "跨境电商示例：传统模式下选品—比价—物料—上架需一周；"
        "OPT「一人跨境电商」方案通过「选品雷达—物料制作—卖点测试」AI 运营系统，"
        "可将核心环节压缩到一个下午——这正是 OPT 对生产力的直接兑现。",
        indent=True,
    )

    # ── 四、六层安全 ──
    add_heading(doc, "四、六层安全体系：企业 AI 守门员", 1)
    add_para(
        doc,
        f"{BRAND} 的安全体系不是功能清单，而是产品骨架。"
        "以下六层自下而上构建信任，任何一层不可绕过：",
        indent=True,
    )
    layers = [
        (
            "第一层：双层规则体系",
            "基础安全规则（30+ 内置 YAML 规则 + Shell 混淆检测）是 AI 行为的最高优先级底线，"
            "任何 Prompt 或 Skill 都无法突破；企业可在此基础上自定义规则、"
            "禁用/启用内置规则、配置 STRICT/SMART/AUTO/OFF 四级审批策略，适配行业合规。",
        ),
        (
            "第二层：统一身份认证",
            "JWT + RBAC 三表权限模型（用户—角色—权限）；"
            "Agent 有效权限取「平台授权」与「提问人身份」的交集——"
            "即使 Agent 理论上可访问某数据，提问人无权限时仍不返回。"
            "频道 allowlist 防止未授权 IM 用户触发 Agent。",
        ),
        (
            "第三层：专属沙箱隔离",
            "容器级 Execution Sandbox（Local Path Jail / Docker Per-Call / Session Container），"
            "遵循最小权限原则：默认 network=none、资源限额、fail_closed；"
            "即使某个 Skill 存在漏洞，攻击面也被严格限制在沙箱边界内。",
        ),
        (
            "第四层：Skill 生态安全",
            "所有 Skill 在创建/启用/导入前自动安全扫描（block/warn/off）；"
            "内容哈希白名单版本锁定；运行时 Tool Guard Policy 引擎对每一步工具调用实时评估；"
            "HIGH/CRITICAL 发现可强制 requires_sandbox。",
        ),
        (
            "第五层：全链路审计",
            "沙箱命令日志、Tool Guard 拦截记录、登录与权限变更、"
            "Skill 扫描告警均可追溯——谁在什么时间、用什么身份、对什么资源做了什么操作。"
            "企业级 SIEM 对接已在路线图（ToB 二次开发）中规划。",
        ),
        (
            "第六层：AI 原生文件系统",
            "从零构建的工作区文件体系：dialog 持久化、memory 分层、"
            "backup/restore 一键回退、per-user 目录隔离；"
            "File Guard 保护 .secret 密钥目录与管理员定义的敏感路径；"
            "任何文件操作在 Path Jail 边界内执行，改错可回退至任意备份版本。",
        ),
    ]
    for title, desc in layers:
        add_heading(doc, title, 3)
        add_para(doc, desc, indent=True)
    add_image(doc, img_paths.get("security"), IMAGES["security"][1])

    # ── 五、架构 ──
    add_heading(doc, "五、产品整体架构", 1)
    add_para(
        doc,
        f"{BRAND} 采用「安全内核 + OPT 编排 + 能力扩展」三层总体架构，"
        "一眼可见产品在企业数字体系中的位置：",
        indent=True,
    )
    add_table(
        doc,
        ["层级", "模块", "职责"],
        [
            ["应用层", "Console · 频道 Gateway · REST API · CLI", "统一人机交互与系统集成入口"],
            ["OPT 编排层", "OrgBuilder · Multi-Agent · Cron/Heartbeat", "岗位人机流程、多 Agent 调度、定时自动化"],
            ["能力层", "Skills · MCP · Builtin Tools · ACP", "行业技能包、工具协议、外部编码 Agent"],
            ["安全内核", "Tool Guard · File Guard · Skill Scanner · Sandbox", "六层安全体系，策略不可绕过"],
            ["数据层", "Workspace · Memory · Secret Store · JWT/MySQL/Redis", "文件、记忆、密钥、身份与权限"],
        ],
    )
    add_para(doc, "请求全链路：", bold=True)
    add_para(
        doc,
        "IM/API 接入 → 身份认证（权限交集）→ Agent/OPT 路由 → Skill 预扫描 "
        "→ Tool Guard Policy →（可选）人工审批 → Sandbox 内执行 "
        "→ 记忆/文件持久化 → 审计留痕 → 结果回传。",
        indent=True,
    )

    # ── 六、功能详解 ──
    add_heading(doc, "六、产品功能详解（图文并茂）", 1)

    features = [
        (
            "6.1 控制台对话 — 企业生产力中枢",
            "统一的 Web 对话界面，支持多模态（图片/音频/视频/文件）、"
            "执行环境切换（沙箱开/关）、工具调用可视化。"
            "员工无需学习 Prompt 工程，即可在权限边界内完成复杂任务。",
            "console_chat",
        ),
        (
            "6.2 安全中心 — 企业 AI 守门员",
            "可视化配置工具守卫、文件防护、技能扫描器、执行沙箱、"
            "Shell 混淆检测、免认证主机白名单；策略修改即时生效，"
            "安全团队可独立完成策略运营，无需改代码。",
            "security",
        ),
        (
            "6.3 技能管理 — 能力市场 + 供应链安全",
            "本地技能池、多 Hub 导入、启用前扫描；"
            "面向 ISV 的开发—审核—上架—分发链路，"
            "目标打造 ToB Skill 能力市场，兼容部分开源 Skill 体系。",
            "skills",
        ),
        (
            "6.4 多智能体 — 按岗位部署 AI 军团",
            "同一实例并行运行多个 Agent，各自独立配置/记忆/技能/频道；"
            "支持 QA、文档、运维等专岗 Agent 协作，"
            "复杂任务异步委派，子会话隔离防污染。",
            "agents",
        ),
        (
            "6.5 频道接入 — 嵌入现有 IM 工作流",
            "钉钉、飞书、企业微信、QQ、Discord、Telegram 等 20+ 频道；"
            "员工在惯用 IM 中与 Agent 对话，零学习成本，"
            "企业无需替换现有协作工具。",
            "channels",
        ),
        (
            "6.6 定时任务与心跳 — 7×24 自动化",
            "Cron 定时任务：日报推送、巡检、数据同步；"
            "Heartbeat 自检：按间隔执行预设问题并回传频道；"
            "Daemon 长周期任务支持，让 Agent 从「被动问答」变为「主动运营」。",
            "cron",
        ),
        (
            "6.7 工作区文件 — AI 原生文件系统",
            "Agent 产出文件统一管理工作区；支持备份/恢复/版本回退；"
            "per-user 目录隔离；File Guard 保护敏感路径——"
            "让 AI 操作文件像 Git 一样可追溯、可回滚。",
            "files",
        ),
        (
            "6.8 模型管理 — 大小模型协同降本",
            "云端 API（OpenAI 兼容）+ 本地 Ollama/LM Studio/llama.cpp；"
            "全局 QPM 流控防 429；敏感数据走本地模型，"
            "复杂推理走云端大模型，兼顾隐私与能力。",
            "models",
        ),
    ]
    for title, desc, img_key in features:
        add_heading(doc, title, 2)
        add_para(doc, desc, indent=True)
        if img_key in IMAGES:
            add_image(doc, img_paths.get(img_key), IMAGES[img_key][1])

    # ── 七、对比 ──
    add_heading(doc, "七、相对 OpenClaw / Hermes 的企业级优势", 1)
    add_para(
        doc,
        "OpenClaw 与 Hermes 代表了 To-C 个人 Agent 的两极："
        "前者以 Gateway + ClawHub 生态见长；后者以自学习记忆与 Skill 自动生成著称。"
        "二者在个人效率场景表现优异，但企业采购时往往卡在安全、治理与合规。",
        indent=True,
    )
    add_table(
        doc,
        ["维度", "OpenClaw", "Hermes", f"{BRAND}（企业 OPT）"],
        [
            ["定位", "个人 AI Gateway", "自学习个人 Agent", "企业级 OPT 平台"],
            ["安全哲学", "能力优先，安全可选", "个人服务器信任", "安全原生，策略不可绕过"],
            ["Shell 沙箱", "可选 Docker", "无企业级沙箱", "三层沙箱 + fail_closed"],
            ["规则引擎", "DM 配对 + allowlist", "无", "30+ 规则 + 自定义 + 审批流"],
            ["Skill 安全", "ClawHub + VirusTotal", "自生成，供应链难控", "上架扫描 + Policy + 白名单"],
            ["身份权限", "个人", "个人", "JWT RBAC + 权限交集"],
            ["文件隔离", "弱", "弱", "Path Jail + File Guard + 备份回退"],
            ["OPT 岗位", "无", "无", "销售/运营/HR/研发等岗位编排"],
            ["审计", "基础", "基础", "多层留痕 + 企业 SIEM 路线"],
            ["部署", "npm/VPS", "curl 一键", "Docker 私有化 + 离线模型"],
        ],
    )
    add_para(
        doc,
        "选型建议：个人极客选 OpenClaw/Hermes；"
        f"企业组织、合规行业、OPT 生产力场景选 {BRAND}。",
        bold=True,
        indent=True,
    )

    # ── 八、行业标杆案例 ──
    add_heading(doc, "八、行业标杆案例", 1)
    add_para(
        doc,
        f"以下案例来自 {BRAND} 在汽车制造、通信运营商、大型央企等行业的实际落地，"
        "覆盖销售、运营、人力、内训等多条 OPT 岗位线。"
        "案例展示的不是「单点聊天机器人」，而是「多 Agent 协同 + 安全可控 + 业务流程嵌入」"
        "的企业级 OPT 交付模式。",
        indent=True,
    )
    add_table(
        doc,
        ["客户", "行业/部门", "OPT 岗位", "核心助手", "业务成效"],
        [
            [
                "北汽极狐",
                "汽车制造 / 门店销售",
                "销售",
                "汽车获客助手、销售数据分析助手",
                "门店自动获客，销售数据实时洞察",
            ],
            [
                "广东移动",
                "通信 / 内训部门",
                "培训与办公",
                "AI 自动化办公助手",
                "内训事务自动化，办公效率提升",
            ],
            [
                "某大型央企 A",
                "集团人力资源部",
                "HR 全生态",
                "招聘/入职/KPI/员工关系助手",
                "人力资源全链路 AI 覆盖",
            ],
            [
                "某大型央企 B",
                "SaaS 平台 / 运营部门",
                "售前·售中·售后",
                "获客/线索/讲解/合同/报告助手",
                "客户全生命周期 OPT 运营",
            ],
        ],
    )

    add_case_block(
        doc,
        "案例一：北汽极狐 — 汽车销售 OPT",
        [
            ("客户", "北汽极狐（ARCFOX）"),
            ("行业", "新能源汽车制造"),
            ("OPT 岗位", "销售（门店一线 + 销售管理）"),
            ("部署形态", "私有化部署 + 门店/总部多 Agent"),
        ],
        "传统汽车销售依赖门店顾问人工跟进线索、手工整理报表，获客响应慢、"
        "数据分散在各门店系统，总部难以及时掌握区域销售态势。"
        "极狐需要一套既能辅助门店自动获客、又能支撑销售数据分析的智能体体系。",
        [
            "【汽车获客助手】嵌入门店销售流程，自动识别高意向客户、"
            "生成跟进话术与试驾邀约，辅助顾问完成线索培育与转化；"
            "支持企微/钉钉等频道触达，门店顾问在 IM 中即可调用。",
            "【销售数据分析助手】汇聚各门店销售、客流、转化数据，"
            "自动生成日报/周报/区域对比分析，帮助销售经理快速定位短板、"
            "调配资源，实现「数据驱动」而非「经验驱动」的管理决策。",
        ],
        [
            "门店获客响应时间显著缩短，顾问可将精力集中在高价值客户面谈",
            "总部销售管理可视化，区域对比与趋势分析从天级压缩到小时级",
            "一人一店即可拥有「获客专员 + 数据分析师」的 OPT 能力组合",
        ],
        [
            "多 Agent 工作区：获客助手与数据分析助手独立配置、独立记忆",
            "频道接入：门店企微/钉钉 IM 零改造接入",
            "Tool Guard + 沙箱：销售脚本与数据处理在策略边界内执行",
            "定时任务：日报/周报自动生成推送",
        ],
    )

    add_case_block(
        doc,
        "案例二：广东移动 — 内训部门 AI 自动化办公",
        [
            ("客户", "广东移动"),
            ("行业", "通信运营商"),
            ("OPT 岗位", "内训与行政办公"),
            ("部署形态", "内网私有化 + 多用户 RBAC"),
        ],
        "内训部门日常承担课程设计、培训组织、材料编写、学员管理等大量事务性工作，"
        "重复性文档处理占用讲师与教务人员大量时间，"
        "且培训资料涉及内部制度与业务知识，不宜使用公有云 AI 服务。",
        [
            "【AI 自动化办公助手】覆盖培训计划起草、课件摘要、"
            "会议纪要整理、学员问答、培训效果初评等场景；",
            "内训师专注课程设计与现场授课，AI 承担材料准备与流程跟进。",
            "【知识库问答】接入内部培训制度与业务文档，"
            "学员与教务人员可通过 IM 或控制台快速检索制度与流程。",
            "【定时任务】培训提醒、作业催交、结业通知自动推送。",
        ],
        [
            "内训材料准备时间大幅压缩，讲师人效显著提升",
            "培训资料与对话数据不出内网，满足运营商数据安全要求",
            "一名内训师 + AI 助手 ≈ 传统小团队的产出能力",
        ],
        [
            "JWT 多用户 RBAC：按部门/角色分配访问权限",
            "File Guard：保护内部培训资料与敏感文档路径",
            "本地模型 + 云端大模型协同：敏感内容本地处理",
            "工作区文件：课件与纪要版本可追溯、可备份",
        ],
    )

    add_case_block(
        doc,
        "案例三：某大型央企 — 人力资源全生态 OPT",
        [
            ("客户", "某大型央企（匿名）"),
            ("行业", "中央企业管理"),
            ("OPT 岗位", "人力资源全生态"),
            ("部署形态", "集团私有化 + 部门级 Agent 矩阵"),
        ],
        "央企人力资源管理覆盖招聘、入职、绩效、员工关系等全生命周期，"
        "环节多、制度严、合规要求高，传统 HR 系统以流程记录为主，"
        "缺乏智能化辅助，事务性咨询占用 HR 大量工时。",
        [
            "【招聘助手】简历智能筛选、岗位匹配评分、面试安排协调、"
            "录用流程提醒，缩短从收简历到发 Offer 的周期。",
            "【入职助手】新员工入职清单自动生成、材料核验提醒、"
            "制度与福利问答、首日/首周引导流程自动化。",
            "【KPI 助手】绩效目标分解建议、考核数据汇总、"
            "异常指标预警、绩效面谈材料初稿生成。",
            "【员工关系助手】制度政策问答、休假/福利/异动流程指引、"
            "员工诉求分类与升级路径建议，减轻 HRBP 重复咨询压力。",
        ],
        [
            "人力资源全链路 AI 覆盖，从「单点工具」升级为「HR 助手矩阵」",
            "事务性 HR 咨询自动化率显著提升，HRBP 聚焦组织发展与员工关怀",
            "权限交集机制确保员工仅可查询本人权限范围内的信息与政策",
            "全链路操作留痕，满足央企内审与合规检查要求",
        ],
        [
            "OrgBuilder 组织架构：HR 各子职能映射独立 Agent",
            "六层安全体系：制度文档 File Guard + 员工数据隔离",
            "JWT RBAC：HR/经理/员工分级权限",
            "Skill 扫描：HR 相关 Skill 上架前安全审计",
        ],
    )

    add_heading(doc, "案例四：某大型央企 — SaaS 平台运营 OPT（售前·售中·售后）", 2)
    add_table(
        doc,
        ["项目要素", "说明"],
        [
            ("客户", "某大型央企（匿名）"),
            ("行业", "SaaS 平台 / 数字化服务"),
            ("OPT 岗位", "运营部门（服务运营 + 销售支持）"),
            ("部署形态", "三阶段渐进式 OPT 落地"),
        ],
    )
    add_para(doc, "业务背景", bold=True)
    add_para(
        doc,
        "该央企运营一套面向政企客户的 SaaS 平台，"
        "售前获客、售中交付、售后运营分属不同团队，"
        "客户全生命周期缺乏统一的 AI 协同体系，"
        "重复性讲解、合同审查、使用报告撰写消耗大量人力。",
        indent=True,
    )
    add_heading(doc, "阶段一：售前 — 服务运营部门", 3)
    add_para(
        doc,
        "目标：扩大有效线索池，提升销售获客效率，"
        "让运营人员从「手工找线索」转向「审核 AI 挖掘结果」。",
        indent=True,
    )
    for item in [
        "【销售获客助手】基于行业标签与客户画像，自动生成触达策略、"
        "营销话术与跟进计划，辅助销售团队扩大有效客户池。",
        "【客户线索挖掘助手】从公开信息、存量客户、行业数据库中"
        "智能挖掘潜在线索，去重、评分、优先级排序，"
        "每日推送高价值线索清单至销售与运营 IM。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "阶段二：售中 — 交付与签约", 3)
    add_para(
        doc,
        "目标：加速客户理解与合同流程，缩短从意向到签约的周期。",
        indent=True,
    )
    for item in [
        "【系统讲解助手】根据客户行业与角色，自动生成 SaaS 平台"
        "功能讲解脚本、演示路径与 FAQ，销售/售前工程师"
        "「一人完成传统售前团队」的讲解准备。",
        "【合同法规助手】对接合同模板与法规知识库，"
        "自动审查合同条款风险、标注不合规项、"
        "生成修订建议，法务人员专注终审而非逐条初审。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "阶段三：售后 — 客户成功与运营", 3)
    add_para(
        doc,
        "目标：持续洞察客户使用情况，自动生成运营报告，"
        "提升续约率与客户满意度。",
        indent=True,
    )
    for item in [
        "【SaaS 平台使用分析总结助手】汇聚用户登录、功能使用、"
        "活跃度、异常行为等数据，自动生成使用分析摘要，"
        "客户成功经理快速掌握客户健康度。",
        "【用户使用季度总结报告助手】按季度自动聚合"
        "各租户/各模块使用数据，生成标准化季度汇报材料，"
        "含趋势图描述、亮点提炼、风险预警与改进建议。",
        "【报告自动生成助手】支持日报/周报/月报/季报模板化输出，"
        "一键推送至管理层 IM 或邮件，运营人员从「写报告」"
        "解放为「审报告、做决策」。",
    ]:
        add_bullet(doc, item)

    add_para(doc, "三阶段整体价值", bold=True)
    for item in [
        "客户全生命周期 OPT 覆盖：售前获客 → 售中交付 → 售后运营",
        "同一 SaaS 平台运营团队，人力不变的情况下服务客户数倍增",
        "各阶段 Agent 独立工作区 + 协作 Skill，数据与记忆不串扰",
        "报告类输出统一走工作区文件系统，版本可追溯、可回退",
    ]:
        add_bullet(doc, item)

    add_para(doc, f"启用的 {BRAND} 能力", bold=True)
    for item in [
        "OPT OrgBuilder：三阶段流程可视化编排",
        "多 Agent 矩阵：10+ 专岗助手并行运行",
        "Cron 定时任务：线索日报、季度报告自动生成",
        "频道接入：销售/运营/法务/管理层 IM 统一触达",
        "Tool Guard + File Guard：合同与客户数据分级保护",
    ]:
        add_bullet(doc, item)

    add_image(doc, img_paths.get("agents"), IMAGES["agents"][1])
    add_image(doc, img_paths.get("cron"), IMAGES["cron"][1])

    add_heading(doc, "8.1 案例共性：从项目到方法论", 2)
    add_table(
        doc,
        ["落地模式", "说明", "适用客户"],
        [
            ["单岗位单助手", "如极狐获客助手", "业务场景单一、需快速见效"],
            ["单部门多助手", "如央企 HR 四助手矩阵", "职能线全链路覆盖"],
            ["三阶段渐进 OPT", "如 SaaS 售前售中售后", "客户生命周期长的平台型业务"],
            ["安全先行部署", "内网私有化 + RBAC + 沙箱", "央企/运营商/金融"],
        ],
    )

    # ── 九、企业效益 ──
    add_heading(doc, "九、企业效益：生产力、成本、合规", 1)
    add_table(
        doc,
        ["效益维度", "具体价值", "可量化参考"],
        [
            ["人效提升", "OPT 一人团队，事务性工作自动化", "岗位产出 +30%~60%"],
            ["成本降低", "本地模型 + 流控，减少 SaaS Token 支出", "推理成本 -40%~70%"],
            ["上市加速", "选品/内容/代码/文档环节压缩", "核心流程从天级到小时级"],
            ["风险可控", "六层安全 + 沙箱，避免「养虾」事故", "高危操作拦截率 99%+"],
            ["合规可证", "权限交集 + 审计留痕 + 私有化", "通过等保/内审概率提升"],
            ["组织弹性", "万人集团到一人公司均可部署", "同一平台横向扩展"],
        ],
    )
    add_para(
        doc,
        "AI 的竞争，已经从模型参数转向「谁能真正嵌入业务流程」。"
        f"{BRAND} 的价值不在于聊天多流畅，而在于"
        "能否让销售、运营、人力、研发等岗位在安全边界内"
        "真正少用人、少出错、少踩雷。",
        indent=True,
    )

    # ── 十、适用人群 ──
    add_heading(doc, "十、适用人群与部署建议", 1)
    personas = [
        ("一人公司与超级个体", "OPT 方案量身打造：选品、内容、数据分析自动化，专注核心业务"),
        ("中小企业与创业团队", "Docker/脚本开箱即用，管理员配置权限与 Skill，团队效率倍增"),
        ("大型组织与跨国公司", "六层安全 + JWT RBAC + 全链路审计，满足数据安全与合规审计"),
        ("开发者与 ISV", "Skill 能力市场：开发—审核—上架—分发，共享 AI 红利"),
    ]
    for title, desc in personas:
        add_heading(doc, title, 3)
        add_para(doc, desc, indent=True)

    add_para(doc, "部署建议：", bold=True)
    for step in [
        "POC：单机 Docker，启用 Local 沙箱 + Tool Guard SMART 模式",
        "试点：接入飞书/钉钉，选 1~2 个 OPT 岗位（如运营+HR）",
        "推广：JWT 多用户 + Docker 沙箱 fail_closed + Skill Scanner block",
        "规模化：部门 OrgBuilder 映射 + 备份策略 + 审计导出",
    ]:
        add_bullet(doc, step)

    # ── 十一、总结 ──
    add_heading(doc, "十一、总结", 1)
    add_para(
        doc,
        f"从北汽极狐的门店销售 OPT，到广东移动的内训自动化，"
        f"再到央企 HR 全生态与 SaaS 三阶段运营，"
        f"{BRAND} 已在汽车、通信、央企等多行业证明："
        "企业需要的不是「更大的聊天窗口」，而是「可投产、可治理、可度量」"
        "的 OPT 智能体平台。",
        indent=True,
    )
    add_para(
        doc,
        f"{BRAND} 不是「带安全功能的个人助手」，而是「架构在安全体系之上的企业 OPT 平台」。"
        "在 OpenClaw、Hermes 点燃 To-C Agent 热潮的同时，"
        "企业需要的不是更大的技能市场，而是可证明、可治理、可隔离的智能体操作系统。"
        f"{BRAND} 以六层安全为骨架、以 OPT 一人团队为触角、"
        "以私有化部署为底座，帮助企业在 AI 时代构建真正可投产的人机协作生产力。",
        indent=True,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("—— 感谢阅读 ——")
    set_run_font(run, size=12, color=(0, 102, 153))
    doc.add_paragraph()
    add_para(doc, f"版权声明：本文档仅供 {BRAND} 产品市场推广使用。产品能力以实际交付版本为准。")

    return doc


def main():
    DOCS.mkdir(parents=True, exist_ok=True)
    print("Downloading console screenshots...")
    img_paths = download_images()
    print(f"Downloaded {len(img_paths)} images.")
    doc = build_document(img_paths)
    doc.save(str(OUTPUT))
    doc.save(str(OUTPUT_EN))
    print(f"Generated: {OUTPUT}")
    print(f"Generated: {OUTPUT_EN}")


if __name__ == "__main__":
    main()
