#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Export ~1000 lines of ai-work-os source code for software copyright filing."""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "src" / "aiwork"
DOCS = ROOT / "docs"
OUTPUT = DOCS / "ai-work-os_软著源代码_1000行.docx"

SOFTWARE_NAME = "ai-work-os 企业级智能体操作系统"
SOFTWARE_VERSION = "V1.0"
TARGET_LINES = 1000

# Core proprietary modules (no qwenpaw branding, no secret/config leakage)
SOURCE_FILES: list[tuple[Path, int | None]] = [
    (SRC / "security" / "sandbox" / "path_jail.py", None),
    (SRC / "security" / "tool_guard" / "engine.py", None),
    (SRC / "security" / "sandbox" / "docker_runner.py", None),
    (SRC / "security" / "skill_scanner" / "scanner.py", None),
    (SRC / "security" / "sandbox" / "session_container_manager.py", None),
    (SRC / "app" / "multi_agent_manager.py", None),
]

BLOCKED_LINE_PATTERNS = [
    re.compile(r"qwenpaw", re.I),
    re.compile(r"copaw", re.I),
    re.compile(r"(?i)(password|passwd|api_key|secret|token)\s*=\s*['\"][^'\"]+['\"]"),
    re.compile(r"(?i)(mysql\+aiomysql|redis://)[^\s'\"]+"),
    re.compile(r"(?i)Bearer\s+[A-Za-z0-9._\-]+"),
    re.compile(r"(?i)(jwt_secret|private_key)\s*=\s*"),
]


def is_blocked_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped or stripped.startswith("#"):
        return False
    return any(p.search(line) for p in BLOCKED_LINE_PATTERNS)


def sanitize_line(line: str) -> str:
    if is_blocked_line(line):
        return "# [redacted for copyright export]"
    return line.rstrip("\n\r")


def collect_source_lines() -> list[tuple[str, str]]:
    """Return list of (filename, line_content)."""
    collected: list[tuple[str, str]] = []
    for path, max_lines in SOURCE_FILES:
        if not path.is_file():
            continue
        rel = path.relative_to(ROOT).as_posix()
        raw_lines = path.read_text(encoding="utf-8").splitlines()
        if max_lines is not None:
            raw_lines = raw_lines[:max_lines]
        collected.append((rel, f"# ===== File: {rel} ====="))
        for line in raw_lines:
            collected.append((rel, sanitize_line(line)))
            if len(collected) >= TARGET_LINES + 50:
                break
        if len(collected) >= TARGET_LINES + 50:
            break
    return collected[:TARGET_LINES]


def set_run_font(run, name: str = "Consolas", size: int = 9):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)


def build_document(lines: list[tuple[str, str]]) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Cover info for 软著
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(SOFTWARE_NAME)
    set_run_font(r, "微软雅黑", 16)
    r.bold = True

    for text in [
        f"软件版本：{SOFTWARE_VERSION}",
        f"源代码摘录（前 {TARGET_LINES} 行）",
        f"生成日期：{date.today().strftime('%Y-%m-%d')}",
        "说明：本文档用于软件著作权登记，已剔除第三方品牌标识及敏感配置信息。",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, "微软雅黑", 11)

    doc.add_page_break()

    # Source code body — one paragraph per line keeps line numbers stable
    line_no = 1
    for _, content in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(f"{line_no:4d}  {content}")
        set_run_font(run)
        line_no += 1

    # Footer summary
    doc.add_page_break()
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = summary.add_run(f"—— 共 {len(lines)} 行源代码 ——")
    set_run_font(run, "微软雅黑", 11)

    modules = sorted({fname for fname, _ in lines if fname.endswith(".py")})
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("涉及源文件模块：")
    set_run_font(run, "微软雅黑", 11)
    run.bold = True
    for mod in modules:
        bp = doc.add_paragraph(style="List Bullet")
        r = bp.add_run(mod)
        set_run_font(r, "微软雅黑", 10)

    return doc


def main():
    DOCS.mkdir(parents=True, exist_ok=True)
    lines = collect_source_lines()
    if len(lines) < TARGET_LINES:
        raise SystemExit(
            f"Only collected {len(lines)} lines, expected {TARGET_LINES}",
        )
    doc = build_document(lines)
    doc.save(str(OUTPUT))
    print(f"Generated: {OUTPUT}")
    print(f"Total lines: {len(lines)}")


if __name__ == "__main__":
    main()
