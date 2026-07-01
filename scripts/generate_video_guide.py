#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate ai-work-os official website video recording guide (Word)."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BRAND = "ai-work-os"
ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
OUTPUT = DOCS / f"{BRAND}_官网展示视频录屏方案.docx"
OUTPUT_EN = DOCS / f"{BRAND}_website_video_recording_guide.docx"


def set_run_font(run, name: str = "微软雅黑", size: int = 11, bold: bool = False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size={1: 18, 2: 14, 3: 12}.get(level, 11), bold=True, color=(0, 51, 102))


def add_para(doc: Document, text: str, bold: bool = False, indent: bool = False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True, size=10)
    for row_data in rows:
        row = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row[i].text = cell_text
            for p in row[i].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10)
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # 封面
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(BRAND)
    set_run_font(run, size=28, bold=True, color=(0, 51, 102))
    for text, size in [
        ("官网展示视频 · 录屏方案", 18),
        ("平台能力 + 行业场景 · 分镜脚本 · 执行清单", 13),
        (f"版本 1.0  |  {date.today().strftime('%Y年%m月')}", 11),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=size, color=(80, 80, 80) if size == 11 else None)
    doc.add_page_break()

    # 目录
    add_heading(doc, "目录", 1)
    for item in [
        "一、核心结论：录什么、怎么讲",
        "二、官网视频矩阵（建议做多条）",
        "三、主视频结构（2 分 30 秒）",
        "四、完整分镜脚本（含旁白）",
        "五、控制台功能录屏优先级",
        "六、行业场景 Agent 录屏清单",
        "七、录屏技术规范",
        "八、录前准备与 Demo 环境清单",
        "九、后期制作要点",
        "十、执行排期建议",
    ]:
        add_para(doc, item)
    doc.add_page_break()

    # 一、核心结论
    add_heading(doc, "一、核心结论：录什么、怎么讲", 1)
    add_para(
        doc,
        "官网主视频建议采用「平台能力 60% + 场景智能体 40%」的组合，"
        "不要只录一个 Agent 聊天。ai-work-os 的差异化是企业 OPT 平台 + 安全原生架构，"
        "若仅展示对话界面，观众会将其与 OpenClaw、Hermes 等 To-C 个人助手混为一谈。",
        indent=True,
    )
    add_table(
        doc,
        ["类型", "占比建议", "录什么", "给谁看"],
        [
            ["平台级", "60%", "工作台、组织架构、安全中心、多 Agent、频道、定时任务", "IT / 安全 / 管理层"],
            ["智能体级", "40%", "1~2 个行业场景 Agent 完成真实任务", "业务负责人 / 采购决策人"],
            ["不建议主视觉", "—", "纯聊天、Debug、Token 统计、模型下载", "内部运维向"],
        ],
    )
    add_para(doc, "叙事顺序（黄金法则）：", bold=True)
    for item in [
        "先用智能体 15 秒展示业务结果（「看，它帮我完成了 XX」）",
        "再用平台 60 秒建立信任（「这是企业级 OPT 平台，不是个人玩具」）",
        "用安全 20 秒做差异化（「OpenClaw/Hermes 没有的治理能力」）",
        "最后 5 秒 CTA（预约演示 / 联系我们）",
    ]:
        add_bullet(doc, item)

    # 二、视频矩阵
    add_heading(doc, "二、官网视频矩阵（建议做多条）", 1)
    add_table(
        doc,
        ["视频名称", "时长", "用途", "首页位置"],
        [
            ["品牌主视频", "2~3 min", "完整展示平台 + 1 个场景", "首屏 Hero 下方 / 关于我们"],
            ["30 秒电梯版", "30s", "仅钩子 + 定位 + CTA", "首屏 Hero 背景自动播放"],
            ["安全专题片", "60s", "六层安全 + 沙箱 + 拦截", "安全/产品页"],
            ["OPT 专题片", "60s", "岗位工作台 + 组织架构", "解决方案页"],
            ["行业案例·极狐销售", "30s", "获客 + 数据分析", "案例页 / 汽车"],
            ["行业案例·央企 HR", "30s", "招聘 + 入职 + KPI", "案例页 / 人力"],
            ["行业案例·SaaS 运营", "45s", "售前售中售后三阶段", "案例页 / 运营"],
            ["行业案例·广东移动", "30s", "内训自动化办公", "案例页 / 通信"],
        ],
    )

    # 三、主视频结构
    add_heading(doc, "三、主视频结构（2 分 30 秒）", 1)
    add_table(
        doc,
        ["时间轴", "画面", "模块", "目的"],
        [
            ["0:00–0:15", "销售日报自动生成完成", "场景 Agent", "钩子：有业务结果"],
            ["0:15–0:30", "岗位工作台总览", "OPT 平台", "一人团队定位"],
            ["0:30–0:50", "组织架构 + OPT 流程", "OrgBuilder", "企业级部署"],
            ["0:50–1:05", "多智能体管理切换", "Agents", "多 Agent 并行"],
            ["1:05–1:25", "安全中心四 Tab", "Security", "安全原生差异化"],
            ["1:25–1:40", "对话中沙箱/执行环境切换", "Chat", "可控执行"],
            ["1:40–1:55", "飞书/钉钉频道接入", "Channels", "嵌入现有 IM"],
            ["1:55–2:10", "定时任务 + 心跳", "Cron/Heartbeat", "7×24 自动化"],
            ["2:10–2:30", "Logo + Slogan + CTA", "品牌", "转化"],
        ],
    )

    # 四、分镜脚本
    add_heading(doc, "四、完整分镜脚本（含旁白）", 1)
    add_para(doc, "主视频：2 分 30 秒版本", bold=True)
    storyboard = [
        ("0:00–0:05", "黑屏 → Logo 淡入", "ai-work-os — 企业级智能体 OPT 平台", "—"),
        ("0:05–0:15", "聊天界面：输入「生成本周门店销售分析」→ 报告输出", "Chat / 销售数据分析助手", "当销售还在手工做报表，你的团队已经拿到了 AI 生成的区域对比分析。"),
        ("0:15–0:25", "岗位工作台：销售、运营、HR 岗位卡片", "Workbench", "ai-work-os 不是聊天工具，是 OPT 平台——One Person Team，一人即一支队伍。"),
        ("0:25–0:35", "组织架构图展开，OPT 协同流程步骤", "OrgChart / OrgBuilder", "从 CEO 到一线门店，每个岗位都可以定义人机分工与 AI 赋能深度。"),
        ("0:35–0:45", "智能体管理：多个 Agent 卡片，切换 QA/销售/HR", "Agents", "销售、人力、运维——每个岗位独立 Agent，独立记忆，互不干扰。"),
        ("0:45–0:55", "安全中心：工具守卫 Tab，规则列表滚动", "Security / ToolGuard", "我们不是给 Agent 贴安全补丁——ai-work-os 架构在安全体系之上。"),
        ("0:55–1:05", "文件防护 + 技能扫描 + 执行沙箱 Tab 快速切换", "Security", "六层纵深安全：规则、身份、沙箱、Skill 供应链、审计、文件系统。"),
        ("1:05–1:15", "聊天界面：执行环境选择器 → 选择 Docker 沙箱", "Chat / Sandbox", "容器级沙箱隔离，最小权限，每一次工具调用都在策略边界内。"),
        ("1:15–1:25", "工具调用被拦截 → 审批弹窗（可选）", "Chat / ToolGuard", "危险操作？自动拦截。高风险？人工审批。企业 AI，必须可控。"),
        ("1:25–1:35", "频道管理：飞书、钉钉、企微卡片启用", "Channels", "不改变员工习惯——在飞书、钉钉里直接调用 AI 能力。"),
        ("1:35–1:45", "定时任务列表 + 新建任务", "CronJobs", "7×24 自动化：日报推送、线索挖掘、季度报告，无需人工值守。"),
        ("1:45–1:55", "技能管理：启用 Skill + 扫描通过提示", "Skills", "Skill 上架前自动安全扫描，运行时 Policy 逐步评估。"),
        ("1:55–2:10", "快速蒙太奇：案例 Logo（极狐/移动/央企）", "案例", "北汽极狐、广东移动、大型央企——已在多行业落地。"),
        ("2:10–2:30", "Logo + Slogan + 按钮", "CTA", "ai-work-os — 架构在安全之上的企业 OPT 平台。预约演示。"),
    ]
    add_table(doc, ["时间", "画面", "页面/模块", "旁白文案"], storyboard)

    add_heading(doc, "4.1 30 秒电梯版脚本", 2)
    add_table(
        doc,
        ["时间", "画面", "旁白"],
        [
            ["0:00–0:08", "销售报告自动生成", "AI 帮你完成业务，不是陪你聊天。"],
            ["0:08–0:18", "安全中心 + 沙箱", "架构在安全之上的企业 OPT 平台。"],
            ["0:18–0:25", "组织架构 + 多 Agent", "一人团队，覆盖销售、人力、运营全岗位。"],
            ["0:25–0:30", "Logo + CTA", "ai-work-os。预约演示。"],
        ],
    )

    # 五、功能优先级
    add_heading(doc, "五、控制台功能录屏优先级", 1)
    add_heading(doc, "5.1 第一优先级（主视频必录）", 2)
    add_table(
        doc,
        ["路由", "页面名称", "录什么", "时长建议"],
        [
            ["/workbench", "岗位工作台", "岗位卡片、进入 Agent", "10s"],
            ["/org-chart", "AI 数字化看板", "组织架构 + OPT 流程", "15s"],
            ["/chat", "聊天", "真实任务 + 工具调用 + 沙箱切换", "25s"],
            ["/security", "安全", "四 Tab 切换概览", "20s"],
            ["/agents", "智能体管理", "多 Agent 列表与切换", "10s"],
        ],
    )
    add_heading(doc, "5.2 第二优先级（补充片 / 分视频）", 2)
    add_table(
        doc,
        ["路由", "页面名称", "录什么", "适用视频"],
        [
            ["/channels", "频道", "飞书/钉钉/企微启用", "集成专题"],
            ["/cron-jobs", "定时任务", "新建 + 列表", "自动化专题"],
            ["/heartbeat", "心跳", "自检配置", "自动化专题"],
            ["/skills", "技能", "启用 + 扫描", "Skill 专题"],
            ["/skill-pool", "Skill Pool", "技能池浏览", "能力市场"],
            ["/tools", "工具", "内置工具开关", "开发者向"],
            ["/mcp", "MCP", "客户端配置", "集成向"],
            ["/models", "模型", "本地+云端（勿录下载）", "技术向"],
            ["/ai-okr", "AI-OKR", "考核看板", "管理向"],
        ],
    )
    add_heading(doc, "5.3 不建议出现在主视频", 2)
    for item in [
        "Debug 调试页、Token 消耗统计",
        "环境变量、备份恢复（偏运维）",
        "模型下载等待过程",
        "登录失败、工具报错、审批卡住等异常",
        "含真实 API Key、密码、客户合同的画面",
    ]:
        add_bullet(doc, item)

    # 六、行业场景
    add_heading(doc, "六、行业场景 Agent 录屏清单", 1)
    cases = [
        (
            "案例一：北汽极狐 — 销售 OPT",
            [
                ("Agent", "汽车获客助手 + 销售数据分析助手"),
                ("输入", "「分析本周华北区门店客流与转化率」"),
                ("过程", "工具调用 → 数据汇总 → 图表描述"),
                ("输出", "结构化销售分析报告（PDF/表格）"),
                ("时长", "30~45 秒"),
            ],
        ),
        (
            "案例二：广东移动 — 内训自动化",
            [
                ("Agent", "AI 自动化办公助手"),
                ("输入", "「整理本次培训课件摘要并生成学员通知」"),
                ("过程", "文档读取 → 摘要 → 定时提醒配置"),
                ("输出", "培训摘要 + 通知草稿"),
                ("时长", "30 秒"),
            ],
        ),
        (
            "案例三：央企 — HR 全生态",
            [
                ("Agent", "招聘 / 入职 / KPI / 员工关系 四助手"),
                ("输入", "分别演示：简历筛选、入职清单、绩效汇总、制度问答"),
                ("过程", "四个 Agent 快速切换（蒙太奇）"),
                ("输出", "各助手产出物截图"),
                ("时长", "45 秒（每条 10s）"),
            ],
        ),
        (
            "案例四：央企 — SaaS 运营三阶段",
            [
                ("阶段一", "销售获客助手 + 线索挖掘助手（15s）"),
                ("阶段二", "系统讲解助手 + 合同法规助手（15s）"),
                ("阶段三", "使用分析 + 季度报告 + 报告自动生成（15s）"),
                ("总时长", "45 秒"),
            ],
        ),
    ]
    for title, items in cases:
        add_heading(doc, title, 2)
        add_table(doc, ["要素", "内容"], items)

    add_para(doc, "每个场景 Agent 录屏公式：", bold=True)
    add_para(doc, "输入（用户问题）→ 过程（工具/Skill 调用可见）→ 输出（可交付物）", indent=True)

    # 七、技术规范
    add_heading(doc, "七、录屏技术规范", 1)
    add_table(
        doc,
        ["项目", "建议值", "说明"],
        [
            ["分辨率", "1920×1080", "官网标准"],
            ["帧率", "30 fps", "UI 演示足够"],
            ["格式", "MP4 (H.264)", "兼容性最好"],
            ["浏览器", "Chrome 最新版", "全屏，缩放 100% 或 110%"],
            ["录屏工具", "OBS / Camtasia / Screen Studio", "OBS 免费够用"],
            ["音频", "48kHz 单声道", "旁白后期配"],
            ["比特率", "8~12 Mbps", "文字清晰不糊"],
        ],
    )
    add_para(doc, "录屏操作要点：", bold=True)
    for item in [
        "关闭系统通知、书签栏、无关浏览器标签",
        "使用专用 Demo 账号，避免真实客户数据",
        "每个镜头录 3 遍，选最流畅的一遍",
        "鼠标移动慢、点击停顿 1 秒，便于后期放大",
        "控制台 UI 统一为 ai-work-os 品牌（去除 QwenPaw 字样）",
    ]:
        add_bullet(doc, item)

    # 八、Demo 清单
    add_heading(doc, "八、录前准备与 Demo 环境清单", 1)
    add_heading(doc, "8.1 环境准备", 2)
    checklist_env = [
        "部署完整 ai-work-os 实例（Docker 推荐），确保控制台可访问",
        "配置至少 1 个云端模型 + 可选 1 个本地模型",
        "预置 4 个 Agent：销售分析、HR 招聘、内训办公、SaaS 运营",
        "每个 Agent 绑定对应 Skill，提前测试对话成功",
        "安全中心四项全部启用（Tool Guard / File Guard / Skill Scanner / Sandbox）",
        "频道至少启用 Console + 1 个 IM（飞书或钉钉演示用）",
        "预置 2 条定时任务（展示用，设为「已执行」状态）",
        "组织架构 Demo 数据（OrgBuilder 示例树）",
    ]
    for item in checklist_env:
        add_bullet(doc, item)

    add_heading(doc, "8.2 假数据准备", 2)
    add_table(
        doc,
        ["数据类型", "示例", "注意"],
        [
            ["客户名", "某汽车集团 / 某央企 / 广东移动", "不用真实全称除非已授权"],
            ["销售数据", "华北区 5 门店虚构 CSV", "数字合理即可"],
            ["简历", "3 份匿名假简历", "脱敏"],
            ["合同", "虚构 SaaS 服务合同", "无真实条款"],
            ["报告", "预生成 1 份销售周报 PDF", "可快速展示输出"],
        ],
    )

    add_heading(doc, "8.3 录屏当日流程", 2)
    for i, step in enumerate([
        "启动服务，清空浏览器缓存，登录 Demo 账号",
        "按分镜脚本顺序逐段录制（先录平台，再录场景）",
        "每段录 3 遍，当场标记最佳 take",
        "录完后检查有无敏感信息、UI 品牌不一致",
        "素材命名：aiworkos_主视频_01_钩子_v2.mp4",
        "交付后期：原始录屏 + 分镜脚本 + 旁白文案",
    ], 1):
        add_bullet(doc, f"{i}. {step}")

    # 九、后期
    add_heading(doc, "九、后期制作要点", 1)
    add_table(
        doc,
        ["环节", "要求"],
        [
            ["字幕", "全程中文字幕，关键词可加粗"],
            ["标注", "关键按钮放大 + 高亮圈（安全、沙箱、OPT）"],
            ["章节", "平台模块切换加章节标题卡"],
            ["音乐", "轻背景音乐，不抢旁白"],
            ["品牌", "片头片尾统一 ai-work-os Logo 与色板"],
            ["导出", "主视频 1080p；电梯版额外导出 720p 竖版（可选）"],
        ],
    )

    # 十、排期
    add_heading(doc, "十、执行排期建议", 1)
    add_table(
        doc,
        ["阶段", "工作内容", "工期", "产出"],
        [
            ["D1", "Demo 环境搭建 + 假数据 + Agent 配置", "1 天", "可演示环境"],
            ["D2", "分镜彩排 + UI 品牌检查", "0.5 天", "彩排录像"],
            ["D3", "正式录屏（主视频 + 4 案例短片）", "1 天", "原始素材"],
            ["D4–D5", "剪辑 + 旁白 + 字幕 + 标注", "2 天", "成片"],
            ["D6", "官网嵌入 + 压缩优化", "0.5 天", "上线"],
        ],
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("—— 方案结束 ——")
    set_run_font(run, size=12, color=(0, 102, 153))
    doc.add_paragraph()
    add_para(doc, f"本文档为 {BRAND} 官网展示视频录屏执行方案，供市场与产品团队使用。")

    return doc


def main():
    DOCS.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(str(OUTPUT))
    doc.save(str(OUTPUT_EN))
    print(f"Generated: {OUTPUT}")
    print(f"Generated: {OUTPUT_EN}")


if __name__ == "__main__":
    main()
